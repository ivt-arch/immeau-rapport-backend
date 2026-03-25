"""
Backend de génération de rapport ITV - Immeau
Reçoit les données JSON de l'app Flutter, génère le .docx et l'envoie par mail.
"""

from flask import Flask, request, jsonify
import re
import os
import io
import base64
import threading
import requests as req_lib
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
import copy

# Regex pour détecter les en-têtes majeurs de section (I –, II –, ..., VIII –, etc.)
_MAJOR_SECTION_RE = re.compile(r'^([IVX]+|\d+)\s+\u2013\s+')

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 80 * 1024 * 1024  # 80 MB

BREVO_API_KEY  = os.environ.get("BREVO_API_KEY", "")
MAIL_FROM      = os.environ.get("MAIL_FROM", "ivt@immeau.fr")
MAIL_TO        = os.environ.get("MAIL_TO",   "ivt@immeau.fr")
CLAUDE_API_KEY = os.environ.get("CLAUDE_API_KEY", "")

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "template.docx")

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


# ─────────────────────────────────────────────
# Helpers python-docx
# ─────────────────────────────────────────────

def _build_adresse_full(adresse: str, cp: str, ville: str) -> str:
    """Construit l'adresse complète sans doublon (si cp/ville déjà dans adresse)."""
    a = adresse.strip()
    parts = [a]
    if cp and cp.strip() and cp.strip() not in a:
        parts.append(cp.strip())
    if ville and ville.strip() and ville.strip().lower() not in a.lower():
        parts.append(ville.strip())
    # Première partie = adresse, reste = CP + ville séparés par une virgule
    if len(parts) == 1:
        return parts[0]
    return parts[0] + ", " + " ".join(parts[1:])


def _strip_sdc_prefix(client: str) -> str:
    """Supprime le préfixe 'SDC DU ' ou 'SDC du ' du nom client si présent
    (le template a déjà 'SDC DU' en texte statique)."""
    return re.sub(r'^SDC\s+DU\s+', '', client, flags=re.IGNORECASE).strip()


def replace_text_in_paragraph(paragraph, old: str, new: str) -> bool:
    """Remplace old par new dans un paragraphe en préservant le formatage du premier run."""
    full = "".join(run.text for run in paragraph.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    for i, run in enumerate(paragraph.runs):
        run.text = new_full if i == 0 else ""
    return True


_FIELD_TAGS = frozenset([
    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar',
    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText',
])


def replace_text_in_paragraph_safe(paragraph, old: str, new: str) -> bool:
    """Remplace old par new dans un paragraphe en ne touchant que les runs
    sans éléments de champ (fldChar, instrText). Préserve les champs PAGE/DATE.

    Important : les runs qui sont des valeurs en cache de champ (positionnés entre
    un fldChar 'separate' et un fldChar 'end') sont AUSSI exclus — sinon leur texte
    (ex: '2' pour le champ PAGE) serait fusionné dans le premier run et s'afficherait
    en double à côté du champ recalculé.

    Note technique : on garde des références dures (liste d'objets) plutôt que des
    id() entiers, car les proxies lxml peuvent être ramassés par le GC entre les deux
    boucles, provoquant des collisions d'adresse mémoire et des faux-positifs.
    """
    _Wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    # Repérer les runs de valeur en cache (entre separate et end).
    # On conserve les références directes (hard refs) pour éviter que le GC lxml
    # les recycle et crée de faux id() identiques sur d'autres éléments.
    cached_run_elems: list = []
    _after_sep = False
    for child in paragraph._element:
        ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if ctag == 'r':
            _fcs = child.findall(f'{{{_Wns}}}fldChar')
            if _fcs:
                _fc_type = _fcs[0].get(f'{{{_Wns}}}fldCharType')
                if _fc_type == 'begin':
                    _after_sep = False
                elif _fc_type == 'separate':
                    _after_sep = True
                elif _fc_type == 'end':
                    _after_sep = False
            elif _after_sep:
                cached_run_elems.append(child)  # référence dure → pas de GC

    # Identifier les runs "purs" : pas de fldChar/instrText ET pas en cache de champ
    safe_runs = [
        run for run in paragraph.runs
        if not any(child.tag in _FIELD_TAGS for child in run._r)
        and not any(run._r is cr for cr in cached_run_elems)  # comparaison par identité
    ]
    full = "".join(run.text for run in safe_runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    for i, run in enumerate(safe_runs):
        run.text = new_full if i == 0 else ""
    return True


def replace_in_doc(doc: Document, replacements: dict):
    """Applique tous les remplacements dans le document entier
    (paragraphes + tables + en-têtes/pieds de page)."""
    # Corps du document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        replace_text_in_paragraph(para, old, str(new))
    for para in doc.paragraphs:
        for old, new in replacements.items():
            replace_text_in_paragraph(para, old, str(new))

    # En-têtes et pieds de page : utilise la version sécurisée pour préserver les champs
    for section in doc.sections:
        for hf in [section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer]:
            if hf is None:
                continue
            for para in hf.paragraphs:
                for old, new in replacements.items():
                    replace_text_in_paragraph_safe(para, old, str(new))
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for old, new in replacements.items():
                                replace_text_in_paragraph_safe(para, old, str(new))


def _get_para_text(elem) -> str:
    return "".join(t.text or "" for t in elem.iter() if t.tag.endswith("}t"))


def _has_numpr(elem) -> bool:
    return any(e.tag.endswith("}numPr") for e in elem.iter())


def _has_underline(elem) -> bool:
    return any(e.tag.endswith("}u") for e in elem.iter())


def _convert_numpr_to_manual_bullet(para, font_size=11):
    """
    Convertit un paragraphe de style liste Word (w:numPr) en puce manuelle •
    identique au style des titres de canalisations insérés dynamiquement :
    • [texte en gras souligné, Arial 11, noir]
    Supprime aussi l'indentation héritée du style liste.
    """
    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    # Récupérer le texte brut avant de vider les runs
    text = para.text.strip()

    # Supprimer tous les runs existants
    for r_elem in list(para._element.findall(f'{{{_W}}}r')):
        para._element.remove(r_elem)

    # Nettoyer les propriétés de liste dans pPr
    pPr = para._element.find(f'{{{_W}}}pPr')
    if pPr is not None:
        for tag_name in ('numPr', 'ind'):
            child = pPr.find(f'{{{_W}}}{tag_name}')
            if child is not None:
                pPr.remove(child)

    # Run puce (non souligné)
    bullet_run = para.add_run('\u2022  ')
    bullet_run.bold = True
    bullet_run.underline = False
    bullet_run.font.name = 'Arial'
    bullet_run.font.size = Pt(font_size)
    bullet_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Run texte (gras + souligné)
    text_run = para.add_run(text)
    text_run.bold = True
    text_run.underline = True
    text_run.font.name = 'Arial'
    text_run.font.size = Pt(font_size)
    text_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _remove_elements(body, elements):
    for elem in elements:
        try:
            # Utiliser getparent() pour être sûr de supprimer depuis le vrai parent
            # (body.remove() échoue silencieusement si l'élément a été déplacé)
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
            else:
                body.remove(elem)
        except Exception:
            pass


def delete_elements_by_text_range(doc: Document, start_contains: str, end_contains: str,
                                   include_start=True, include_end=True):
    body = doc.element.body
    body_children = list(body)

    in_range = False
    to_remove = []
    found_start = False

    for elem in body_children:
        text = _get_para_text(elem) if elem.tag.endswith("}p") else ""

        if not found_start and start_contains in text:
            found_start = True
            in_range = True
            if include_start:
                to_remove.append(elem)
            continue

        if in_range:
            if end_contains and end_contains in text:
                if include_end:
                    to_remove.append(elem)
                break
            else:
                to_remove.append(elem)

    _remove_elements(body, to_remove)
    return len(to_remove)


def delete_single_paragraph(doc: Document, contains: str):
    body = doc.element.body
    for elem in list(body):
        if elem.tag.endswith("}p"):
            if contains in _get_para_text(elem):
                try:
                    body.remove(elem)
                except ValueError:
                    pass
                return True
    return False


def delete_section_by_title(doc: Document, section_title: str):
    """Supprime une section de conclusion identifiée par son titre."""
    body = doc.element.body
    paras = list(body)

    target_idx = None
    for i, elem in enumerate(paras):
        if elem.tag.endswith("}p") and _has_numpr(elem) and _has_underline(elem):
            text = _get_para_text(elem)
            if section_title in text:
                target_idx = i
                break
    if target_idx is None:
        return

    to_remove = [paras[target_idx]]
    for i in range(target_idx + 1, len(paras)):
        elem = paras[i]
        if elem.tag.endswith("}p"):
            text = _get_para_text(elem)
            if _MAJOR_SECTION_RE.match(text):
                break
        if elem.tag.endswith("}p") and _has_numpr(elem) and _has_underline(elem):
            break
        to_remove.append(elem)

    _remove_elements(body, to_remove)


def insert_paragraphs_before(doc: Document, anchor_contains: str, paragraphs_data: list,
                             ref_list_para=None):
    """
    Insère des paragraphes avant l'élément contenant anchor_contains.
    paragraphs_data : liste de dicts { 'text', 'bold', 'underline', 'bullet', 'font_size' }
    ref_list_para   : paragraphe de référence dont on copie le style liste Word (numPr + ind)
                      pour les entrées avec 'bullet': True. Si None, utilise une puce manuelle •.
    """
    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body
    anchor = None
    for elem in body:
        if elem.tag.endswith("}p") and anchor_contains in _get_para_text(elem):
            anchor = elem
            break
    if anchor is None:
        return

    # Extraire numPr et ind du paragraphe de référence (une seule fois)
    ref_numPr = None
    ref_ind   = None
    if ref_list_para is not None:
        ref_pPr = ref_list_para._element.find(f'{{{_W}}}pPr')
        if ref_pPr is not None:
            _np = ref_pPr.find(f'{{{_W}}}numPr')
            if _np is not None:
                ref_numPr = copy.deepcopy(_np)
            _ind = ref_pPr.find(f'{{{_W}}}ind')
            if _ind is not None:
                ref_ind = copy.deepcopy(_ind)

    last_inserted = None
    for pdata in paragraphs_data:
        new_para = doc.add_paragraph()
        text = pdata.get('text', '').strip()
        font_size = pdata.get('font_size', 11)
        is_bold = bool(pdata.get('bold'))
        is_underline = bool(pdata.get('underline'))

        if pdata.get('bullet'):
            if ref_numPr is not None:
                # Copier le style liste Word du paragraphe de référence
                pPr = new_para._element.find(f'{{{_W}}}pPr')
                if pPr is None:
                    pPr = etree.SubElement(new_para._element, f'{{{_W}}}pPr')
                    new_para._element.insert(0, pPr)
                # Ajouter numPr (et ind si présent) dans pPr
                existing_np = pPr.find(f'{{{_W}}}numPr')
                if existing_np is not None:
                    pPr.remove(existing_np)
                pPr.append(copy.deepcopy(ref_numPr))
                if ref_ind is not None:
                    existing_ind = pPr.find(f'{{{_W}}}ind')
                    if existing_ind is not None:
                        pPr.remove(existing_ind)
                    pPr.append(copy.deepcopy(ref_ind))
            else:
                # Fallback : puce manuelle • dans un run séparé sans soulignement
                bullet_run = new_para.add_run("\u2022  ")
                bullet_run.bold = is_bold
                bullet_run.underline = False
                bullet_run.font.name = 'Arial'
                bullet_run.font.size = Pt(font_size)
                bullet_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        run = new_para.add_run(text)
        run.bold = is_bold
        run.underline = is_underline
        # Police Arial 11 systématiquement + couleur noir explicite (évite héritage rouge)
        run.font.name = 'Arial'
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        space_before = pdata.get('space_before', 0)
        space_after  = pdata.get('space_after', 0)
        if space_before:
            new_para.paragraph_format.space_before = Pt(space_before)
        if space_after:
            new_para.paragraph_format.space_after = Pt(space_after)

        new_elem = new_para._element
        body.remove(new_elem)
        if last_inserted is None:
            anchor.addprevious(new_elem)
        else:
            last_inserted.addnext(new_elem)
        last_inserted = new_elem


def _replace_placeholder_with_paragraphs(doc: Document, placeholder: str, text: str):
    """
    Remplace un paragraphe contenant un placeholder par un ou plusieurs paragraphes,
    en splitant sur les doubles sauts de ligne (\n\n).
    Police Arial 11 appliquée à chaque paragraphe inséré.
    """
    # Trouver le paragraphe cible dans le body
    body = doc.element.body
    target_elem = None
    target_para = None
    for para in doc.paragraphs:
        if placeholder in para.text:
            target_para = para
            target_elem = para._element
            break
    if target_para is None:
        return

    # Nettoyer le texte et splitter
    parts = [p.strip() for p in text.split('\n\n') if p.strip()]
    if not parts:
        replace_text_in_paragraph(target_para, placeholder, "")
        return

    # Remplacer le contenu du premier paragraphe
    full = "".join(run.text for run in target_para.runs)
    new_full = full.replace(placeholder, parts[0])
    for i, run in enumerate(target_para.runs):
        run.text = new_full if i == 0 else ""
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    # Insérer les paragraphes supplémentaires
    last_elem = target_elem
    for part in parts[1:]:
        new_para = doc.add_paragraph()
        run = new_para.add_run(part)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        new_elem = new_para._element
        body.remove(new_elem)
        last_elem.addnext(new_elem)
        last_elem = new_elem


def _join_french_list(items: list) -> str:
    """Joint une liste en français : 'a, b, c et d'."""
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    return ", ".join(items[:-1]) + " et " + items[-1]


# ─────────────────────────────────────────────
# Gestion des photos
# ─────────────────────────────────────────────

def _insert_facade_photo(doc: Document, facade_bytes: bytes):
    FACADE_PH = "\u2018ici photo de la fa\u00e7ade en mode portrait, 9,87cm de hauteur\u2019"
    for para in doc.paragraphs:
        if FACADE_PH in para.text:
            for run in para.runs:
                run.text = ""
            try:
                run = para.add_run()
                run.add_picture(io.BytesIO(facade_bytes), height=Cm(9.87))
            except Exception as e:
                print(f"[WARN] Erreur insertion photo façade : {e}")
            return


def _clear_cell_content(cell):
    """Supprime tout le contenu textuel et les images d'une cellule,
    en conservant uniquement les propriétés de paragraphe."""
    for p in cell._tc.findall(f'{{{W_NS}}}p'):
        for child in list(p):
            if not child.tag.endswith('}pPr'):
                p.remove(child)


def _fill_photo_cell(doc: Document, cell, photo_data: dict):
    """Remplace le contenu d'une cellule par une photo + commentaire."""
    image_b64 = photo_data.get('image_base64', '')
    commentaire = photo_data.get('commentaire', '')

    # Vider entièrement la cellule (texte et images)
    _clear_cell_content(cell)

    # Supprimer tous les paragraphes sauf le premier pour éviter les lignes vides
    tc = cell._tc
    all_paras = tc.findall(f'{{{W_NS}}}p')
    for p in all_paras[1:]:
        tc.remove(p)

    # Utiliser le premier (et maintenant unique) paragraphe pour l'image
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        image_bytes = base64.b64decode(image_b64)
        run = para.add_run()
        run.add_picture(io.BytesIO(image_bytes), height=Cm(9.87))
    except Exception as e:
        print(f"[WARN] Erreur insertion photo : {e}")
        para.add_run("Photo non disponible")

    # Commentaire directement dans le paragraphe suivant (sans ligne vide intermédiaire)
    if commentaire:
        comment_para = cell.add_paragraph()
        comment_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = comment_para.add_run(commentaire)
        run.bold = True
        run.italic = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        # Désactiver toute transformation de casse (préserver majuscules/minuscules du texte saisi)
        run.font.all_caps = False
        run.font.small_caps = False


def _clear_cell(cell):
    """Vide le contenu d'une cellule."""
    _clear_cell_content(cell)


def _clear_table(tbl):
    """Vide toutes les cellules d'une table."""
    for row in tbl.rows:
        for cell in row.cells:
            _clear_cell_content(cell)


def _fill_photo_tables(doc: Document, photos_commentees: list):
    """
    Remplit les tables de photos du complément photographique.
    Identifie les tables photo dynamiquement par leur contenu (placeholder 'ici photo')
    pour éviter les problèmes de décalage d'indice causés par les tables insérées
    dynamiquement (ex: cadastre placeholder).
    """
    photo_idx = 0
    total_photos = len(photos_commentees)
    body = doc.element.body

    # ── Trouver dynamiquement les tables photo par leur contenu ──────────────
    # Le template a des cellules contenant 'ici photo' dans les tables photo
    photo_tables = []
    for tbl in doc.tables:
        try:
            cell_text = tbl.cell(0, 0).text.lower()
        except Exception:
            cell_text = ''
        if 'ici photo' in cell_text:
            photo_tables.append(tbl)

    if not photo_tables:
        return  # Aucune table photo trouvée

    for tbl_pos, tbl in enumerate(photo_tables):
        rows_to_remove = []

        for row in tbl.rows:
            row_has_photo = False
            n_cols = min(2, len(row.cells))
            for col_idx in range(n_cols):
                cell = row.cells[col_idx]
                if col_idx > 0 and row.cells[col_idx]._tc is row.cells[0]._tc:
                    continue  # cellule fusionnée
                if photo_idx < total_photos:
                    _fill_photo_cell(doc, cell, photos_commentees[photo_idx])
                    photo_idx += 1
                    row_has_photo = True
                else:
                    _clear_cell_content(cell)

            if not row_has_photo:
                rows_to_remove.append(row._tr)

        tbl_elem = tbl._tbl
        for tr in rows_to_remove:
            try:
                tbl_elem.remove(tr)
            except Exception:
                pass

        if photo_idx >= total_photos:
            # Construire la liste des tables à supprimer :
            # - la table courante si elle est vide (total_photos=0 ou fin exacte sur une table)
            # - toutes les tables photo restantes
            _W_TR = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'
            tables_to_delete = []
            current_tbl_elem = tbl._tbl
            if len(current_tbl_elem.findall(_W_TR)) == 0:
                # Table courante entièrement vidée → la supprimer aussi
                tables_to_delete.append(tbl)
            tables_to_delete.extend(photo_tables[tbl_pos + 1:])

            for del_tbl in tables_to_delete:
                del_elem = del_tbl._tbl
                # Supprimer les paragraphes vides/sauts de page qui précèdent
                prev_elem = del_elem.getprevious()
                while (prev_elem is not None
                       and prev_elem.tag.endswith('}p')
                       and not _get_para_text(prev_elem).strip()):
                    prev_to_del = prev_elem
                    prev_elem = prev_elem.getprevious()
                    try:
                        prev_to_del.getparent().remove(prev_to_del)
                    except Exception:
                        pass
                try:
                    del_elem.getparent().remove(del_elem)
                except Exception:
                    pass
            break


# ─────────────────────────────────────────────
# Logique principale de remplissage du template
# ─────────────────────────────────────────────


def _add_cadastre_placeholder(doc: Document):
    """
    Encadre la section Description du site dans un tableau 2 colonnes.
    Colonne gauche : texte de description. Colonne droite : encadré jaune cadastre.
    """
    from docx.oxml.ns import qn
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    body = doc.element.body

    # Trouver les limites : premier paragraphe après "II – DESCRIPTION DU SITE"
    # et jusqu'au paragraphe "III – GÉOLOGIE IN SITU" (non inclus)
    desc_start_elem = None
    desc_end_before_elem = None
    found_heading = False

    for elem in list(body):
        if not elem.tag.endswith('}p'):
            if found_heading and desc_start_elem is not None:
                # table inside description section, skip
                continue
            continue
        text = _get_para_text(elem)
        if 'DESCRIPTION DU SITE' in text.upper() and not found_heading:
            found_heading = True
            continue
        if found_heading:
            if desc_start_elem is None:
                desc_start_elem = elem
            if ('G\u00c9OLOGIE' in text.upper() or 'III' in text or
                    'CANALISATIONS' in text.upper() or
                    _MAJOR_SECTION_RE.match(text)):
                desc_end_before_elem = elem
                break

    if desc_start_elem is None:
        return

    # Collecter tous les paragraphes de description
    collecting = False
    desc_elems = []
    for elem in list(body):
        if elem is desc_start_elem:
            collecting = True
        if collecting:
            if desc_end_before_elem is not None and elem is desc_end_before_elem:
                break
            desc_elems.append(elem)

    if not desc_elems:
        return

    # Créer un tableau 2 colonnes, 1 ligne, sans bordures visibles
    tbl = doc.add_table(rows=1, cols=2)
    tbl_elem = tbl._tbl

    # Supprimer le style de table par défaut (pour éviter les bordures)
    tbl_pr = tbl_elem.find(f'{{{W_NS}}}tblPr')
    if tbl_pr is None:
        tbl_pr = etree.SubElement(tbl_elem, f'{{{W_NS}}}tblPr')
    # Forcer la largeur totale
    tbl_w = tbl_pr.find(f'{{{W_NS}}}tblW')
    if tbl_w is None:
        tbl_w = etree.SubElement(tbl_pr, f'{{{W_NS}}}tblW')
    tbl_w.set(f'{{{W_NS}}}w', '9356')  # largeur en twips (~16.5cm page A4 avec marges)
    tbl_w.set(f'{{{W_NS}}}type', 'dxa')
    # Supprimer les bordures de table
    tbl_borders = tbl_pr.find(f'{{{W_NS}}}tblBorders')
    if tbl_borders is None:
        tbl_borders = etree.SubElement(tbl_pr, f'{{{W_NS}}}tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border_el = tbl_borders.find(f'{{{W_NS}}}{side}')
        if border_el is None:
            border_el = etree.SubElement(tbl_borders, f'{{{W_NS}}}{side}')
        border_el.set(f'{{{W_NS}}}val', 'none')
        border_el.set(f'{{{W_NS}}}sz', '0')
        border_el.set(f'{{{W_NS}}}space', '0')
        border_el.set(f'{{{W_NS}}}color', 'auto')

    # Largeur des colonnes (grille)
    tbl_grid = tbl_elem.find(f'{{{W_NS}}}tblGrid')
    if tbl_grid is None:
        tbl_grid = etree.SubElement(tbl_elem, f'{{{W_NS}}}tblGrid')
    for col_elem in list(tbl_grid):
        tbl_grid.remove(col_elem)
    col_left_w = etree.SubElement(tbl_grid, f'{{{W_NS}}}gridCol')
    col_left_w.set(f'{{{W_NS}}}w', '6237')   # ~11 cm
    col_right_w = etree.SubElement(tbl_grid, f'{{{W_NS}}}gridCol')
    col_right_w.set(f'{{{W_NS}}}w', '3119')  # ~5.5 cm

    row = tbl.rows[0]
    left_cell = row.cells[0]
    right_cell = row.cells[1]

    # Définir les largeurs de cellules
    for cell, w in [(left_cell, '6237'), (right_cell, '3119')]:
        tc = cell._tc
        tc_pr = tc.find(f'{{{W_NS}}}tcPr')
        if tc_pr is None:
            tc_pr = etree.SubElement(tc, f'{{{W_NS}}}tcPr')
        tc_w = tc_pr.find(f'{{{W_NS}}}tcW')
        if tc_w is None:
            tc_w = etree.SubElement(tc_pr, f'{{{W_NS}}}tcW')
        tc_w.set(f'{{{W_NS}}}w', w)
        tc_w.set(f'{{{W_NS}}}type', 'dxa')

    # Déplacer les paragraphes de description dans la cellule gauche
    left_tc = left_cell._tc
    # Supprimer le paragraphe vide initial de la cellule gauche
    for p in list(left_tc.findall(f'{{{W_NS}}}p')):
        left_tc.remove(p)
    for elem in desc_elems:
        try:
            body.remove(elem)
        except Exception:
            pass
        left_tc.append(elem)

    # Cellule droite : fond jaune + texte placeholder cadastre
    right_tc = right_cell._tc
    tc_pr_r = right_tc.find(f'{{{W_NS}}}tcPr')
    if tc_pr_r is None:
        tc_pr_r = etree.SubElement(right_tc, f'{{{W_NS}}}tcPr')
    # Fond jaune
    shd = etree.SubElement(tc_pr_r, f'{{{W_NS}}}shd')
    shd.set(f'{{{W_NS}}}val', 'clear')
    shd.set(f'{{{W_NS}}}color', 'auto')
    shd.set(f'{{{W_NS}}}fill', 'FFF100')
    # Alignement vertical centré
    v_align = etree.SubElement(tc_pr_r, f'{{{W_NS}}}vAlign')
    v_align.set(f'{{{W_NS}}}val', 'center')

    # Texte placeholder
    right_para = right_cell.paragraphs[0] if right_cell.paragraphs else right_cell.add_paragraph()
    right_para.clear()
    right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_r = right_para.add_run("IMPRIM ÉCRAN DU\nCADAST RE\nÀ RÉALISER AU BUREAU")
    # Better: two runs with proper line breaks
    right_para.clear()
    lines = ["IMPRIM ÉCRAN DU CADASTRE", "À RÉALISER AU BUREAU"]
    for idx_l, line in enumerate(lines):
        run_r = right_para.add_run(line)
        run_r.bold = True
        run_r.font.name = 'Arial'
        run_r.font.size = Pt(10)
        run_r.font.color.rgb = RGBColor(0x00, 0x1F, 0x5B)
        if idx_l < len(lines) - 1:
            run_r.add_break()

    # Insérer le tableau à l'emplacement du premier paragraphe de description
    desc_start_elem.addprevious(tbl_elem)


def build_rapport(data: dict) -> bytes:
    doc = Document(TEMPLATE_PATH)

    # ── Extraction des données ──────────────────────────────────────────
    adresse        = data.get("adresseProjet", "")
    ville          = data.get("villeProjet", "")
    cp             = data.get("cpProjet", "")
    client_raw     = data.get("client", "")
    mo_delegue     = data.get("moDelegue", "")
    adresse_mo     = data.get("adresseMoDelegue", "")
    ville_mo       = data.get("villeMoDelegue", "")
    cp_mo          = data.get("cpMoDelegue", "")
    moe            = data.get("moe", "")
    adresse_moe    = data.get("adresseMoe", "")
    ville_moe      = data.get("villeMoe", "")
    cp_moe         = data.get("cpMoe", "")
    devis          = data.get("devis", "")
    date_rapport   = data.get("dateRapport", "")
    redacteur      = data.get("redacteur", "")
    verificateur   = data.get("verificateur", "")
    titre_etude    = data.get("titreEtude", "")
    reglement      = data.get("reglementApplicable", "Ville de Paris")

    # Adresses sans doublon
    adresse_full     = _build_adresse_full(adresse, cp, ville)
    adresse_mo_full  = _build_adresse_full(adresse_mo, cp_mo, ville_mo)
    adresse_moe_full = _build_adresse_full(adresse_moe, cp_moe, ville_moe)

    # Client sans préfixe "SDC DU" (déjà dans le template en texte statique)
    client = _strip_sdc_prefix(client_raw)

    desc_site      = data.get("descriptionSite", "")
    parcelle       = data.get("parcelleCadastre", "")
    section_cad    = data.get("sectionCadastre", "")
    objet_mission  = data.get("objetMission", "")

    paragraphes            = data.get("paragraphesSelectionnes", [])
    reglementations        = data.get("reglementationsSelectionnees", [])

    # BPE
    bpe_present            = data.get("bpePresent", False)
    bpe_type               = data.get("bpeTypeBranchement", "")
    bpe_phrase             = data.get("bpePhraseGeneree", "")

    # Zones canalisations
    batiment_zones         = data.get("batimentZones", [])
    cour_zones             = data.get("courZones", [])

    # Installations sanitaires (liste d'items sélectionnés, ex. ["le siphon de sol en sous-sol"])
    installations_items    = data.get("installationsSanitairesItems", [])

    # Regards de visite non étanches
    regards_noms           = data.get("regardsNonEtanchesNoms", [])    # ex. ["R1","R2","T5"]
    regards_texte          = data.get("regardsNonEtanchesTexte", "")   # conclusion GPT

    # Commentaires réglementaires
    commentaire_colonne_ep     = data.get("commentaireColonneEP", "")
    commentaire_regard_limite  = data.get("commentaireRegardLimite", "regard de visite maçonné")
    fosse_profondeur           = data.get("fosseProfondeur", "")
    fosse_trappe               = data.get("fosseTrappe", False)
    commentaire_separatif      = data.get("commentaireSeparatif", "")
    commentaire_restaurants    = data.get("commentaireRestaurants", "")
    commentaire_garages        = data.get("commentaireGarages", "")

    # Conditions de travaux
    cond_colonne_ep      = data.get("condTravauxColonneEP", False)
    cond_terrassement    = data.get("condTravauxTerrassement", "")
    cond_profondeur      = data.get("condTravauxProfondeur", "")
    cond_sol             = data.get("condTravauxSol", "")
    cond_difficulte      = data.get("condTravauxDifficulte", "")
    cond_pave            = data.get("condTravauxPave", False)

    is_paris = cp.startswith("75") if cp else False

    # ── Nettoyage de objet_mission : supprimer le préambule ───────────────
    # La page objet_mission_page envoie "La présente étude est demandée par X.\n\nCette mission a pour objectif Y."
    # Le template a déjà "La présente étude est demandée par le SDC du [adresse].\nCette mission a pour objectif [placeholder]."
    # On extrait uniquement la partie objectif (après "Cette mission a pour objectif ").
    objet_clean = objet_mission
    m = re.search(r'Cette mission a pour objectif\s+', objet_clean, re.IGNORECASE | re.DOTALL)
    if m:
        objet_clean = objet_clean[m.end():].strip()
    # Supprimer aussi "La présente étude..." si l'objectif n'est pas trouvé
    if not m:
        objet_clean = re.sub(
            r'^La pr[ée]sente [ée]tude est demand[ée]e par[^.]+\.\s*',
            '', objet_clean, flags=re.IGNORECASE
        ).strip()
    # Supprimer les sauts de ligne (causent des espaces larges avec texte justifié)
    objet_clean = objet_clean.replace('\n', ' ').replace('  ', ' ').strip()

    # ── 0a. En-têtes de section : chiffres romains → arabes + texte blanc ──
    # Le style Heading 1 a la couleur 365F91 (bleu) ; on passe à FFFFFF (blanc)
    h1_style = doc.styles['Heading 1']
    rPr = h1_style._element.find(
        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    if rPr is not None:
        color_elem = rPr.find(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
        if color_elem is not None:
            color_elem.set(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val',
                'FFFFFF')
            # Supprimer les attributs de couleur thème pour forcer le blanc
            for attr in list(color_elem.attrib.keys()):
                if 'theme' in attr.lower():
                    del color_elem.attrib[attr]

    # Remplacer chiffres romains par arabes dans headings ET sommaire
    ROMAN_TO_ARABIC = [
        ("VIII \u2013 ", "8 \u2013 "),
        ("VII \u2013 ", "7 \u2013 "),
        ("VI \u2013 ", "6 \u2013 "),
        ("IV \u2013 ", "4 \u2013 "),
        ("III \u2013 ", "3 \u2013 "),
        ("II \u2013 ", "2 \u2013 "),
        ("V \u2013 ", "5 \u2013 "),
        ("I \u2013 ", "1 \u2013 "),
    ]
    for old_r, new_r in ROMAN_TO_ARABIC:
        replace_in_doc(doc, {old_r: new_r})

    # ── 0b-bis. Nettoyer toutes les lignes vides du SOMMAIRE ────────────
    # Le template contient de nombreux paragraphes TM1 vides après les entrées
    # du sommaire (ex : 19 lignes). Avec spacing-before=180 par para, même 2
    # d'entre eux peuvent déborder sur la page suivante selon la longueur du
    # sommaire (rapports avec beaucoup de sections). On les supprime TOUS.
    # On ajoute aussi un saut de page explicite avant le PRÉAMBULE du corps
    # pour garantir une séparation propre entre la page SOMMAIRE et la suite.
    _W_S = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    _body_s = doc.element.body
    in_sommaire_zone = False
    preambule_toc_seen = False
    to_del_toc = []
    preambule_body_elem = None

    for elem in list(_body_s):
        if not elem.tag.endswith('}p'):
            if in_sommaire_zone:
                break  # Une table = fin de zone sommaire
            continue
        text = _get_para_text(elem).strip()
        pPr = elem.find(f'{{{_W_S}}}pPr')
        style_val = ''
        if pPr is not None:
            ps = pPr.find(f'{{{_W_S}}}pStyle')
            if ps is not None:
                style_val = ps.get(f'{{{_W_S}}}val', '')

        if text == 'SOMMAIRE':
            in_sommaire_zone = True
            continue
        if not in_sommaire_zone:
            continue

        if style_val == 'TM1':
            if not text:  # Ligne TM1 vide → supprimer systématiquement
                to_del_toc.append(elem)
            else:
                if text == 'PR\u00c9AMBULE':
                    preambule_toc_seen = True
        elif text == 'PR\u00c9AMBULE' and preambule_toc_seen:
            # Deuxième "PRÉAMBULE" = celui du corps du document
            preambule_body_elem = elem
            break
        elif not text:
            to_del_toc.append(elem)
        else:
            # Premier paragraphe non-vide non-TM1 non-PRÉAMBULE = corps
            preambule_body_elem = elem
            break

    # Supprimer les lignes vides en excès
    for elem in to_del_toc:
        try:
            elem.getparent().remove(elem)
        except Exception:
            pass

    # Ajouter un saut de page avant le PRÉAMBULE du corps (s'il n'en a pas déjà un)
    if preambule_body_elem is not None:
        prev = preambule_body_elem.getprevious()
        prev_has_pb = False
        if prev is not None and prev.tag.endswith('}p'):
            for br in prev.iter(f'{{{_W_S}}}br'):
                if br.get(f'{{{_W_S}}}type', '') == 'page':
                    prev_has_pb = True
                    break
        if not prev_has_pb:
            pb_para = doc.add_paragraph()
            run = pb_para.add_run()
            br_elem = etree.SubElement(run._r, f'{{{_W_S}}}br')
            br_elem.set(f'{{{_W_S}}}type', 'page')
            pb_elem = pb_para._element
            _body_s.remove(pb_elem)
            preambule_body_elem.addprevious(pb_elem)

    # ── 0b. Supprimer P70/P71 AVANT les remplacements globaux ────────────
    # Ces paragraphes contiennent les placeholders 'adresse et 'numéro qui
    # seraient écrasés par l'étape 2 avant de pouvoir être supprimés.
    delete_single_paragraph(doc, "Le site \u00e0 l\u2019\u00e9tude est situ\u00e9 au \u2018adresse")
    delete_single_paragraph(doc, "(parcelle cadastrale \u2018num\u00e9ro")

    # ── 1. Remplacements dans les tables (page de garde) ─────────────────

    # Table 0 : titre avec l'adresse
    tbl0 = doc.tables[0]
    for row in tbl0.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_text_in_paragraph(para,
                    "\u2018ADRESSE DU PROJET, CODE POSTAL VILLE\u2018",
                    adresse_full.upper())

    # Table 1 : acteurs, devis, date, rédacteur, vérificateur
    if len(doc.tables) > 1:
        tbl1 = doc.tables[1]
        cell_mod = tbl1.rows[1].cells[0]
        for para in cell_mod.paragraphs:
            replace_text_in_paragraph(para,
                "\u2018ici Maitre d\u2019ouvrage d\u00e9l\u00e9gu\u00e9\u2019", mo_delegue)
            replace_text_in_paragraph(para,
                "\u2018ici adresse du maitre d\u2019ouvrage d\u00e9l\u00e9gu\u00e9\u2018", adresse_mo)
            replace_text_in_paragraph(para,
                "\u2018ici code postal du MOD\u2019 \u2018ici ville MOD\u2019",
                f"{cp_mo} {ville_mo}")

        cell_moe = tbl1.rows[1].cells[2]
        for para in cell_moe.paragraphs:
            replace_text_in_paragraph(para,
                "\u2018ici Maitre d\u2019oeuvre\u2019", moe)
            replace_text_in_paragraph(para,
                "\u2018ici adresse du maitre d\u2019oeuvre\u2018", adresse_moe)
            replace_text_in_paragraph(para,
                "\u2018ici code postal du MO\u2019 \u2018ici ville MO\u2019",
                f"{cp_moe} {ville_moe}")

        row3 = tbl1.rows[3]
        replace_text_in_paragraph(row3.cells[0].paragraphs[0],
            "\u2018ici num\u00e9ro de devis\u2019", devis)
        replace_text_in_paragraph(row3.cells[1].paragraphs[0],
            "\u2018ici mois de g\u00e9n\u00e9ration du rapport + ann\u00e9e\u2019", date_rapport)
        replace_text_in_paragraph(row3.cells[2].paragraphs[0],
            "\u2018ici r\u00e9dacteur du rapport\u2019", redacteur)
        replace_text_in_paragraph(row3.cells[3].paragraphs[0],
            "\u2018ici v\u00e9rifi\u00e9 par\u2026\u2019", verificateur)

    # ── 2. Remplacements globaux ──────────────────────────────────────────
    replacements = {
        # Page de garde – SDC (le "SDC DU" est statique dans le template)
        "\u2018Client/maitre d\u2019ouvrage\u2019": client.upper(),

        # Titre d'étude (P9) et en-tête de page
        "\u2018ici titre d\u2019\u00e9tude\u2019": titre_etude,

        # Règlement (P52)
        "\u2018remplir ici selon le r\u00e8glement choisi\u2019": reglement,

        # Section I – Objet de la mission
        "\u2018adresse, code postal ville\u2019": adresse_full,

        # Objet de la mission (placeholder dans P68)
        "\u2018ce qui est coch\u00e9 ou tap\u00e9 dans la page objet de la mission\u2019": objet_clean,

        # Section II – Description du site (parcelle)
        "\u2018num\u00e9ro de parcelle cadastrale\u2019": parcelle,
        "\u2018section cadastrale\u2019": section_cad,

        # En-tête de page : 'objet de la prestation'
        "\u2018objet de la prestation\u2019": titre_etude,

        # Pied de page : 'numéro de devis'
        "\u2018num\u00e9ro de devis\u2019": devis,
    }
    replace_in_doc(doc, replacements)

    # ── 3. Supprimer "(si séléctionné dans l'application)" dans la table des matières ─
    INSTR_TDM = " (si s\u00e9l\u00e9ctionn\u00e9 dans l\u2019application)"
    INSTR_TDM2 = " (si s\u00e9l\u00e9ctionn\u00e9 dans l'application)"  # apostrophe ASCII
    for para in doc.paragraphs:
        for marker in [INSTR_TDM, INSTR_TDM2]:
            if marker in para.text:
                replace_text_in_paragraph(para, marker, "")

    # ── 4. Description du site : remplacer le placeholder par plusieurs paragraphes ──
    DESC_PH = "\u2018phrase g\u00e9n\u00e9r\u00e9e avec Claude IA sur la page description du site\u2019"
    if desc_site:
        _replace_placeholder_with_paragraphs(doc, DESC_PH, desc_site)
    else:
        replace_in_doc(doc, {DESC_PH: ""})

    # ── 4b. Supprimer les \n initiaux dans la Section IV (Section IV a des ──
    # paragraphes qui commencent par \n, créant des espaces visuels indésirables)
    replace_in_doc(doc, {
        "\nLe plan de la cour": "Le plan de la cour",
        "\nLe plan du sous-sol": "Le plan du sous-sol",
    })

    # ── 4c. Encadrer la description du site avec un placeholder cadastre ──
    _add_cadastre_placeholder(doc)

    # ── 6. Section IV – Paris vs Hors-Paris ──────────────────────────────
    PARIS_MARKER      = "\u2018si rapport dans paris mettre ce paragraphe\u00a0:\u2019"
    HORS_PARIS_MARKER = "\u2018si rapport en dehors de paris"
    V_CONCLUSIONS     = "5 \u2013 CONCLUSIONS"  # Après remplacement romain→arabe step 0a

    if is_paris:
        # Supprimer le marqueur Paris ET le paragraphe vide qui suit
        delete_elements_by_text_range(doc,
            start_contains=PARIS_MARKER,
            end_contains="Le rapport d\u2019inspection t\u00e9l\u00e9vis\u00e9e",
            include_start=True, include_end=False)
        delete_elements_by_text_range(doc,
            start_contains=HORS_PARIS_MARKER,
            end_contains=V_CONCLUSIONS,
            include_start=True, include_end=False)
    else:
        delete_elements_by_text_range(doc,
            start_contains=PARIS_MARKER,
            end_contains=HORS_PARIS_MARKER,
            include_start=True, include_end=False)
        # Supprimer le marqueur Hors-Paris ET le paragraphe vide qui suit
        delete_elements_by_text_range(doc,
            start_contains=HORS_PARIS_MARKER,
            end_contains="Le rapport d\u2019inspection t\u00e9l\u00e9vis\u00e9e",
            include_start=True, include_end=False)
    # Supprimer le paragraphe vide en bas de section IV (entre dernière ligne légende et section suivante)
    # Ce paragraphe (P96 dans le template) devient un saut de ligne visuel indésirable après la légende
    body_ch = list(doc.element.body)
    for i, elem in enumerate(body_ch):
        if elem.tag.endswith("}p") and "garantissant plus une bonne" in _get_para_text(elem):
            # Supprimer le paragraphe vide suivant s'il existe
            if i + 1 < len(body_ch):
                next_e = body_ch[i + 1]
                if next_e.tag.endswith("}p") and not _get_para_text(next_e).strip():
                    try:
                        doc.element.body.remove(next_e)
                    except Exception:
                        pass
            break

    # ── 7. Section V – BPE ───────────────────────────────────────────────
    BPE_MARKER    = "\u2018si dans la page de l\u2019application Branchement"
    BPE_PHRASE_PH = "\u2018phrase de conclusion en fonction des cases \u00e0 cocher"
    CANALAPP_PH   = "Canalisations apparentes en caves b\u00e2timent"

    if not bpe_present:
        delete_elements_by_text_range(doc,
            start_contains=BPE_MARKER,
            end_contains=CANALAPP_PH,
            include_start=True, include_end=False)
    else:
        for para in doc.paragraphs:
            if "si dans la page de l" in para.text and "BPE" in para.text:
                replace_text_in_paragraph(para,
                    "\u2018si dans la page de l\u2019application Branchement du Particulier \u00e0 l\u2019\u00e9gout on coche pr\u00e9sence d\u2019un BPE (Paris)\u00a0: oui faire apparaitre ce titre et le texte ci dessous\u00a0:\u2019 ",
                    "")
                replace_text_in_paragraph(para,
                    "\u2018si dans la page de l\u2019application Branchement du Particulier \u00e0 l\u2019\u00e9gout on coche pr\u00e9sence d\u2019un BPE (Paris)\u00a0: oui faire apparaitre ce titre et le texte ci dessous\u00a0:\u2019",
                    "")
                break

        FERME_SCHEMA_CAPTION  = "Sch\u00e9ma de principe du branchement particulier ferm\u00e9"
        OUVERT_SCHEMA_CAPTION = "Sch\u00e9ma de principe du branchement particulier ouvert"
        CANAL_SCHEMA_CAPTION  = "Sch\u00e9ma de principe du branchement particulier canalis\u00e9"
        FERME_INSTR  = "\u2018mettre le sch\u00e9ma et la l\u00e9gende ci dessus si dans la page branchement du particulier \u00e0 l\u2019\u00e9gout on coche BPE ferm\u00e9\u2019"
        OUVERT_INSTR = "\u2018mettre le sch\u00e9ma et la l\u00e9gende ci dessus si dans la page branchement du particulier \u00e0 l\u2019\u00e9gout on coche BPE ouvert\u2019"
        CANAL_INSTR  = "\u2018mettre le sch\u00e9ma et la l\u00e9gende ci dessus si dans la page branchement du particulier \u00e0 l\u2019\u00e9gout on coche BPE canalis\u00e9\u2019"

        if bpe_type == 'ferme':
            # Supprimer depuis FERME_INSTR jusqu'à CANAL_INSTR (inclus)
            # Cela supprime les images et légendes des schémas ouvert et canalisé
            # tout en conservant les schémas fermé (P118-P123)
            delete_elements_by_text_range(doc,
                start_contains=FERME_INSTR,
                end_contains=CANAL_INSTR,
                include_start=True, include_end=True)
        elif bpe_type == 'ouvert':
            delete_elements_by_text_range(doc,
                start_contains=FERME_SCHEMA_CAPTION,
                end_contains=FERME_INSTR,
                include_start=True, include_end=True)
            delete_single_paragraph(doc, OUVERT_INSTR)
            delete_elements_by_text_range(doc,
                start_contains=CANAL_SCHEMA_CAPTION,
                end_contains=CANAL_INSTR,
                include_start=True, include_end=True)
        elif bpe_type == 'canalise':
            delete_elements_by_text_range(doc,
                start_contains=FERME_SCHEMA_CAPTION,
                end_contains=FERME_INSTR,
                include_start=True, include_end=True)
            delete_elements_by_text_range(doc,
                start_contains=OUVERT_SCHEMA_CAPTION,
                end_contains=OUVERT_INSTR,
                include_start=True, include_end=True)
            delete_single_paragraph(doc, CANAL_INSTR)
        else:
            delete_elements_by_text_range(doc,
                start_contains=FERME_SCHEMA_CAPTION,
                end_contains=FERME_INSTR,
                include_start=True, include_end=True)
            delete_elements_by_text_range(doc,
                start_contains=OUVERT_SCHEMA_CAPTION,
                end_contains=OUVERT_INSTR,
                include_start=True, include_end=True)
            delete_elements_by_text_range(doc,
                start_contains=CANAL_SCHEMA_CAPTION,
                end_contains=CANAL_INSTR,
                include_start=True, include_end=True)

        for para in doc.paragraphs:
            if BPE_PHRASE_PH.split('\u00e0')[0][1:] in para.text:
                replace_text_in_paragraph(para, para.text, bpe_phrase)
                break

    # ── 8. Section V – Zones de canalisations ───────────────────────────
    CANALAPP_GENERIC = "Canalisations apparentes en caves b\u00e2timent \u2026."
    COLONNE_EP_TITLE = "Colonne d\u2019eaux pluviales de fa\u00e7ade"

    delete_elements_by_text_range(doc,
        start_contains=CANALAPP_GENERIC,
        end_contains=COLONNE_EP_TITLE,
        include_start=True, include_end=False)

    zone_paragraphs = []
    show_app = "Canalisations Apparentes" in paragraphes
    show_ent = "Canalisations Enterrées" in paragraphes

    for zone in batiment_zones:
        zone_name  = zone.get("zoneName", "Bâtiment")
        app_phrase = zone.get("appPhrase", "").strip()
        ent_phrase = zone.get("entPhrase", "").strip()

        if show_app and zone.get("appPresentes") and app_phrase:
            zone_paragraphs.append({
                'text': f"Canalisations apparentes en caves {zone_name}",
                'bold': True, 'underline': True, 'bullet': True,
                'space_before': 12, 'space_after': 4,
            })
            zone_paragraphs.append({
                'text': app_phrase, 'bold': False, 'space_after': 6,
            })

        if show_ent and zone.get("entPresentes") and ent_phrase:
            zone_paragraphs.append({
                'text': f"Canalisations enterr\u00e9es sous {zone_name}",
                'bold': True, 'underline': True, 'bullet': True,
                'space_before': 12, 'space_after': 4,
            })
            zone_paragraphs.append({
                'text': ent_phrase, 'bold': False, 'space_after': 6,
            })

    for zone in cour_zones:
        zone_name  = zone.get("zoneName", "espace extérieur")
        ent_phrase = zone.get("entPhrase", "").strip()

        if show_ent and zone.get("entPresentes") and ent_phrase:
            zone_paragraphs.append({
                'text': f"Canalisations enterr\u00e9es sous espaces ext\u00e9rieurs ({zone_name})",
                'bold': True, 'underline': True, 'bullet': True,
                'space_before': 12, 'space_after': 4,
            })
            zone_paragraphs.append({
                'text': ent_phrase, 'bold': False, 'space_after': 6,
            })

    if zone_paragraphs:
        anchor = None
        for candidate in [COLONNE_EP_TITLE,
                           "Regard de limite de propri\u00e9t\u00e9",
                           "Installations sanitaires en sous-sol",
                           "Ancienne fosse d\u2019aisance",
                           "Regards de visite non \u00e9tanches",
                           "R\u00e9seau s\u00e9paratif",
                           "Cas des eaux us\u00e9es provenant",
                           "Cas des eaux provenant des garages",
                           "VI \u2013 COMPL\u00c9MENT"]:
            for para in doc.paragraphs:
                if candidate in para.text:
                    anchor = candidate
                    break
            if anchor:
                break
        if anchor:
            # Trouver un paragraphe de référence avec le style liste Word du template
            # (numPr + underline) pour que les titres de canalisations aient le même
            # style que "Regard de limite de propriété", "Colonne d'eaux pluviales", etc.
            ref_list_para = None
            for para in doc.paragraphs:
                if _has_numpr(para._element) and _has_underline(para._element):
                    ref_list_para = para
                    break
            insert_paragraphs_before(doc, anchor, zone_paragraphs,
                                     ref_list_para=ref_list_para)

    # ── 9. Sections de conclusions – supprimer les non sélectionnées ─────
    SECTION_MAP = {
        "Colonne d\u2019eaux pluviales de fa\u00e7ade":               "Colonne d\u2019eaux pluviales de fa\u00e7ade",
        "Regard de limite de propri\u00e9t\u00e9":                    "Regard de limite de propri\u00e9t\u00e9",
        "Installations sanitaires en sous-sol":                       "Installations sanitaires en sous-sol",
        "Ancienne fosse d\u2019aisance":                              "Ancienne fosse d\u2019aisance",
        "Regards de visite non \u00e9tanches":                        "Regards de visite non \u00e9tanches",
        "R\u00e9seau s\u00e9paratif":                                 "R\u00e9seau s\u00e9paratif",
        "Cas des eaux us\u00e9es \u2013 Restaurants":                 "Cas des eaux us\u00e9es provenant des restaurants",
        "Cas des eaux \u2013 Garages & Stations de lavage":           "Cas des eaux provenant des garages",
    }

    for flutter_key, template_title in SECTION_MAP.items():
        if flutter_key not in reglementations:
            delete_section_by_title(doc, template_title)
    # Ventilations des r\u00e9seaux : toujours supprim\u00e9e (page retir\u00e9e de l\u2019application)
    delete_section_by_title(doc, "Ventilations des r\u00e9seaux")

    # ── 10. Supprimer instructions résiduelles dans les titres ───────────
    for para in doc.paragraphs:
        if _has_numpr(para._element) and _has_underline(para._element):
            for instr_marker in [
                " \u2018int\u00e9grer ce paragraphe si coch\u00e9 dans selection des paragraphes\u2019",
                "\u2018int\u00e9grer ce paragraphe si coch\u00e9 dans selection des paragraphes\u2019",
            ]:
                if instr_marker in para.text:
                    replace_text_in_paragraph(para, instr_marker, "")
                    break

    # ── 11. Règlement d'assainissement ───────────────────────────────────
    if reglement and "Paris" not in reglement:
        replace_in_doc(doc, {
            "r\u00e8glement d\u2019assainissement de la ville de Paris":
                f"r\u00e8glement d\u2019assainissement de {reglement}",
        })

    # ── 12. Installations sanitaires en sous-sol ─────────────────────────
    # Suppression du texte règlementaire boilerplate
    for _boilerplate_contains in [
        "les hauteurs d\u2019eau dans le r\u00e9seau d\u2019assainissement",
        "L\u2019usager doit se pr\u00e9munir de toutes les cons\u00e9quences de ce fonctionnement du r\u00e9seau, notamment en cas de pr\u00e9sence d\u2019installations sanitaires",
        "De m\u00eame, tous les orifices sur ces canalisations, situ\u00e9s \u00e0 un niveau inf\u00e9rieur",
        "Enfin, tout appareil d\u2019\u00e9vacuation se trouvant \u00e0 un niveau inf\u00e9rieur",
    ]:
        delete_single_paragraph(doc, _boilerplate_contains)
    # Recommandation dynamique ou "À remplir"
    if installations_items:
        items_text = _join_french_list(installations_items)
        replace_in_doc(doc, {
            "pour : x.": f"pour : {items_text}."
        })
    else:
        replace_in_doc(doc, {
            "\nAu vu de la configuration du r\u00e9seau d\u2019assainissement de la construction, la mise en place d\u2019un clapet anti-retour sera n\u00e9cessaire pour : x.":
                "\u00c0 remplir",
            "Au vu de la configuration du r\u00e9seau d\u2019assainissement de la construction, la mise en place d\u2019un clapet anti-retour sera n\u00e9cessaire pour : x.":
                "\u00c0 remplir",
        })

    # ── 13. Regards de visite non étanches ──────────────────────────────
    # Suppression des paragraphes boilerplate règlementaires (dupliqués depuis installations sanitaires)
    for _boilerplate_contains in [
        "les hauteurs d\u2019eau dans le r\u00e9seau d\u2019assainissement, en fonctionnement normal",
        "L\u2019usager doit se pr\u00e9munir de toutes les cons\u00e9quences de ce fonctionnement du r\u00e9seau, notamment en cas de pr\u00e9sence d\u2019installations sanitaires en sous-sol",
        "De m\u00eame, tous les orifices sur ces canalisations, situ\u00e9s \u00e0 un niveau inf\u00e9rieur \u00e0 celui de la voie vers laquelle se fait l\u2019\u00e9vacuation doivent \u00eatre normalement obtures",
    ]:
        delete_single_paragraph(doc, _boilerplate_contains)
    # Recommandation dynamique ou "À remplir"
    if regards_noms or regards_texte:
        if regards_noms:
            noms_text = _join_french_list(regards_noms)
            replace_in_doc(doc, {
                "Les regards de visite x n\u2019\u00e9tant":
                    f"Les regards de visite {noms_text} n\u2019\u00e9tant"
            })
        if regards_texte:
            replace_in_doc(doc, {
                "il conviendra de les reprendre. X.":
                    f"il conviendra de les reprendre. {regards_texte}"
            })
        else:
            replace_in_doc(doc, {"il conviendra de les reprendre. X.": "il conviendra de les reprendre."})
    else:
        replace_in_doc(doc, {
            "Les regards de visite x n\u2019\u00e9tant pas \u00e9tanche (photos ci-dessous), il conviendra de les reprendre. X.":
                "\u00c0 remplir",
        })

    # ── 14. Remplacements règlementaires spécifiques ─────────────────────
    # Colonne EP de façade : remplace le texte statique par le commentaire saisi
    if commentaire_colonne_ep:
        replace_in_doc(doc, {
            "pas de colonne ep": commentaire_colonne_ep,
        })

    # Regard de limite de propriété : suppression du boilerplate règlementaire
    delete_single_paragraph(doc, "les canalisations d\u2019\u00e9vacuation des eaux us\u00e9es et pluviales sont dot\u00e9es d\u2019un regard de visite")
    # Recommandation dynamique ou "À remplir"
    if commentaire_regard_limite and commentaire_regard_limite != "regard de visite ma\u00e7onn\u00e9":
        replace_in_doc(doc, {
            "la mise en place d\u2019un regard de visite ma\u00e7onn\u00e9 en limite de propri\u00e9t\u00e9":
                f"la mise en place d\u2019un {commentaire_regard_limite} en limite de propri\u00e9t\u00e9",
        })
    else:
        replace_in_doc(doc, {
            "Dans le cas pr\u00e9sent, la mise en place d\u2019un regard de visite ma\u00e7onn\u00e9 en limite de propri\u00e9t\u00e9 sera n\u00e9cessaire.":
                "\u00c0 remplir",
        })

    # Ancienne fosse d'aisance : construit la phrase dynamique
    if fosse_profondeur:
        trappe_phrase = " avec une trappe d\u2019acc\u00e8s" if fosse_trappe else ""
        fosse_sentence = (
            f"Une ancienne fosse d\u2019aisance d\u2019une profondeur de {fosse_profondeur}"
            f" m\u00e8tres{trappe_phrase} est pr\u00e9sente."
        )
        replace_in_doc(doc, {
            "Une ancienne fosse d\u2019aisance d\u2019une profondeur de 2 m\u00e8tres avec une trappe d\u2019acc\u00e8s est pr\u00e9sente.":
                fosse_sentence,
            "Une ancienne fosse d\u2019aisance d\u2019une profondeur de 2 m\u00e8tres avec une trappe d\u2019acc\u00e8s est pr\u00e9sente. ":
                fosse_sentence + " ",
        })

    # Réseau séparatif : suppression du boilerplate règlementaire
    for _boilerplate_contains in [
        "vous avez l\u2019obligation de vous assurer que votre logement",
        "En syst\u00e8me s\u00e9paratif, les eaux us\u00e9es et pluviales sont r\u00e9cup\u00e9r\u00e9es",
        "les canalisations et les collecteurs acheminent les eaux us\u00e9es",
        "les collecteurs pluviaux entra\u00eenent les eaux de pluie",
        "la tendance est de limiter au maximum la sollicitation de la canalisation",
        "Citer l\u2019article.",
    ]:
        delete_single_paragraph(doc, _boilerplate_contains)
    # Recommandation dynamique ou "À remplir"
    if commentaire_separatif:
        replace_in_doc(doc, {
            "Le r\u00e8glement d\u2019assainissement x impose, avant le raccordement au domaine public, la mise en place d\u2019un syst\u00e8me s\u00e9paratif d\u2019\u00e9vacuation des eaux us\u00e9es et pluviales.":
                commentaire_separatif,
        })
    else:
        replace_in_doc(doc, {
            "Le r\u00e8glement d\u2019assainissement x impose, avant le raccordement au domaine public, la mise en place d\u2019un syst\u00e8me s\u00e9paratif d\u2019\u00e9vacuation des eaux us\u00e9es et pluviales.":
                "\u00c0 remplir",
        })

    # Restaurants : suppression du boilerplate règlementaire
    for _boilerplate_contains in [
        "Ces \u00e9tablissements sont susceptibles de rejeter des eaux excessivement charg\u00e9es en graisses",
        "Ils doivent \u00eatre \u00e9quip\u00e9s d\u2019un syst\u00e8me de pr\u00e9traitement de leurs effluents",
        "Les \u00e9tablissements (restaurants, cantines",
    ]:
        delete_single_paragraph(doc, _boilerplate_contains)
    # Recommandation dynamique ou "À remplir"
    if commentaire_restaurants:
        replace_in_doc(doc, {
            "il faudra \u00e9quiper le commerce x.":
                f"il conviendra : {commentaire_restaurants}.",
        })
    else:
        replace_in_doc(doc, {
            "Dans notre cas, il faudra \u00e9quiper le commerce x.":
                "\u00c0 remplir",
        })

    # Garages : suppression du boilerplate règlementaire (P212)
    delete_single_paragraph(doc, "Afin de ne pas rejeter dans les \u00e9gouts ou dans les caniveaux, des hydrocarbures")
    # Garages : P213 contient boilerplate + placeholder ; on remplace en 2 passes
    # Passe 1 : supprimer le boilerplate statique de début de P213
    replace_in_doc(doc, {
        "Cette obligation s\u2019applique \u00e9galement aux parcs de stationnement publics et aux parkings d\u2019immeubles, couverts ou non, susceptibles d\u2019accueillir plus d\u00e9couverts ou non, susceptibles.\n\n": "",
        "Cette obligation s\u2019applique \u00e9galement aux parcs de stationnement publics et aux parkings d\u2019immeubles, couverts ou non, susceptibles d\u2019accueillir plus d\u00e9couverts ou non, susceptibles.\n": "",
    })
    # Passe 2 : remplacer le placeholder dynamique
    if commentaire_garages:
        replace_in_doc(doc, {
            "Dans notre cas, il faudra donc x.":
                f"Dans le cas pr\u00e9sent, il conviendra : {commentaire_garages}.",
        })
    else:
        replace_in_doc(doc, {
            "\nDans notre cas, il faudra donc x.": "\u00c0 remplir",
            "Dans notre cas, il faudra donc x.": "\u00c0 remplir",
        })

    # ── 15. Conditions de travaux ─────────────────────────────────────────
    # P230 : présence d'une colonne EP (conditionnel)
    if not cond_colonne_ep:
        delete_single_paragraph(doc,
            "D\u2019autre part, la mise aux normes de la colonne d\u2019eaux pluviales de fa\u00e7ade")

    # P232 : type de terrassement
    TERR_MAP = {
        "manuel_total": "devront se faire manuellement en raison du manque d\u2019acc\u00e8s pour des engins m\u00e9caniques.",
        "mecanique":    "pourront se faire m\u00e9caniquement.",
        "mixte":        "devront se faire manuellement en raison du manque d\u2019acc\u00e8s pour des engins m\u00e9caniques dans certaines zones mais pourront se faire m\u00e9caniquement dans les zones o\u00f9 le ou les engins pourraient circuler.",
    }
    if cond_terrassement and cond_terrassement in TERR_MAP:
        replace_in_doc(doc, {
            "devront se faire manuellement en raison en raison du manque d\u2019acc\u00e8s pour des engins m\u00e9caniques.":
                TERR_MAP[cond_terrassement],
        })

    # P233 : profondeur + sol + difficulté
    if cond_profondeur or cond_sol or cond_difficulte:
        prof_text = cond_profondeur if cond_profondeur else "1"
        sol_text  = cond_sol if cond_sol else "remblais"
        diff_map  = {
            "difficile": "pourraient s\u2019av\u00e9rer difficiles et des purges de blocs anguleux pourraient \u00eatre n\u00e9cessaires.",
            "facile":    "ne devraient pas poser de probl\u00e8mes particuliers.",
        }
        diff_text = diff_map.get(cond_difficulte, "ne devraient pas poser de probl\u00e8mes particuliers.")
        new_p233 = (
            f"La profondeur des fouilles pour la pose des canalisations devrait \u00eatre comprise entre 0 et "
            f"{prof_text} m\u00e8tres de profondeur. Le sol attendu sur place \u00e9tant {sol_text}, "
            f"les terrassements {diff_text}"
        )
        replace_in_doc(doc, {
            "La profondeur des fouilles pour la pose des canalisations devrait \u00eatre comprise entre 0 et 1 m\u00e8tres de profondeur. Le sol attendu sur place \u00e9tant des remblais, les terrassements ne devraient pas poser de probl\u00e8mes particuliers.":
                new_p233,
        })

    # P245-P246 : réfection du pavage (conditionnel)
    if not cond_pave:
        delete_single_paragraph(doc, "R\u00e9fection du pavage")
        delete_single_paragraph(doc, "La r\u00e9fection des autobloquants devra respecter")

    # ── 17. Photo de façade ──────────────────────────────────────────────
    FACADE_PH = "\u2018ici photo de la fa\u00e7ade en mode portrait, 9,87cm de hauteur\u2019"
    photo_facade_b64 = data.get("photoFacade")
    if photo_facade_b64:
        try:
            facade_bytes = base64.b64decode(photo_facade_b64)
            _insert_facade_photo(doc, facade_bytes)
        except Exception as e:
            print(f"[WARN] Erreur décodage photo façade : {e}")
    else:
        for para in doc.paragraphs:
            if FACADE_PH in para.text:
                replace_text_in_paragraph(para, FACADE_PH, "")
                break

    # ── 18. Photos avec commentaires (section VI) ────────────────────────
    photos_commentees = data.get("photosCommentees", [])
    _fill_photo_tables(doc, photos_commentees)

    # ── 19. Forcer la mise à jour des champs (numéros de page) ───────────
    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    # 19a. w:updateFields dans settings → recalcule tous les champs à l'ouverture
    settings_elem = doc.settings.element
    existing_uf = settings_elem.find(f'{{{_W}}}updateFields')
    if existing_uf is None:
        uf = etree.SubElement(settings_elem, f'{{{_W}}}updateFields')
        uf.set(f'{{{_W}}}val', '1')
    else:
        existing_uf.set(f'{{{_W}}}val', '1')
    # 19b. Cibler uniquement le champ PAGE dans les pieds de page :
    # - dirty="1" sur son fldChar begin → recalcul à l'ouverture
    # - Suppression complète du/des run(s) de valeur cachée (entre separate et end)
    #   pour éviter que certains rendeurs mobiles affichent l'ancien "2" + le
    #   numéro recalculé côte-à-côte (ex : "Page 22" au lieu de "Page 2").
    # NB : on ne touche pas aux autres champs du footer (FILENAME, DATE, etc.).
    for section in doc.sections:
        for hf in [section.footer, section.even_page_footer, section.first_page_footer]:
            if hf is None:
                continue
            for para in hf.paragraphs:
                _p = para._element
                # Machine à états sur les <w:r> enfants directs du paragraphe
                state = 'outside'   # outside | in_begin | in_instr | after_sep
                begin_fc = None
                is_page_field = False
                cached_runs = []
                for child in list(_p):
                    ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if ctag != 'r':
                        continue
                    fcs = child.findall(f'{{{_W}}}fldChar')
                    instrs = child.findall(f'{{{_W}}}instrText')
                    if fcs:
                        fc_type = fcs[0].get(f'{{{_W}}}fldCharType')
                        if fc_type == 'begin':
                            state = 'in_begin'
                            begin_fc = fcs[0]
                            is_page_field = False
                            cached_runs = []
                        elif fc_type == 'separate':
                            state = 'after_sep' if is_page_field else 'outside'
                        elif fc_type == 'end':
                            if is_page_field and state == 'after_sep':
                                # Supprimer les runs de valeur cachée
                                for cr in cached_runs:
                                    try:
                                        cr.getparent().remove(cr)
                                    except Exception:
                                        pass
                            state = 'outside'
                            begin_fc = None
                            is_page_field = False
                            cached_runs = []
                    elif instrs:
                        if state == 'in_begin' and 'PAGE' in (instrs[0].text or ''):
                            is_page_field = True
                            state = 'in_instr'
                            if begin_fc is not None:
                                begin_fc.set(f'{{{_W}}}dirty', '1')
                        elif state in ('in_begin', 'in_instr'):
                            state = 'in_instr'
                    elif state == 'after_sep' and is_page_field:
                        cached_runs.append(child)

    # ── 20. Sérialise en mémoire ─────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────
# Rapport généré par Claude IA
# ─────────────────────────────────────────────

def _call_claude_api(system_prompt: str, user_prompt: str, max_tokens: int = 4000) -> str:
    """Appelle l'API Claude Anthropic et retourne le texte."""
    if not CLAUDE_API_KEY:
        raise Exception("CLAUDE_API_KEY manquante dans les variables d'environnement")
    headers = {
        "x-api-key": CLAUDE_API_KEY,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    body = {
        "model": "claude-sonnet-4-5",
        "max_tokens": max_tokens,
        "system": system_prompt,
        "messages": [{"role": "user", "content": user_prompt}],
    }
    resp = req_lib.post("https://api.anthropic.com/v1/messages",
                        headers=headers, json=body, timeout=90)
    if resp.status_code == 200:
        return resp.json()["content"][0]["text"].strip()
    raise Exception(f"Claude API erreur {resp.status_code}: {resp.text[:300]}")


# ═══════════════════════════════════════════════════════════════
# RAPPORT IA — Palette Imméau & helpers XML
# ═══════════════════════════════════════════════════════════════

_IA_NAVY  = '1F4E79'   # Bleu marine principal
_IA_BLUE  = '2E75B6'   # Bleu accent
_IA_LCYAN = 'DEEAF1'   # Bleu très clair (fond bandeau adresse)
_IA_LGRAY = 'F5F5F5'   # Gris clair (fond objet)
_IA_WHITE = 'FFFFFF'
_IA_W_NS  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _ia_rgb(hex6: str) -> RGBColor:
    h = hex6.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _ia_set_cell_bg(cell, hex6: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(f'{{{_IA_W_NS}}}shd'):
        tcPr.remove(old)
    shd = etree.SubElement(tcPr, f'{{{_IA_W_NS}}}shd')
    shd.set(f'{{{_IA_W_NS}}}val', 'clear')
    shd.set(f'{{{_IA_W_NS}}}color', 'auto')
    shd.set(f'{{{_IA_W_NS}}}fill', hex6.upper())


def _ia_get_or_add_tblPr(tbl_elem):
    """Retourne w:tblPr existant ou en crée un en tête de table (compatible python-docx 1.x)."""
    W = _IA_W_NS
    tblPr = tbl_elem.find(f'{{{W}}}tblPr')
    if tblPr is None:
        tblPr = etree.Element(f'{{{W}}}tblPr')
        tbl_elem.insert(0, tblPr)
    return tblPr


def _ia_set_table_no_border(tbl):
    tblPr = _ia_get_or_add_tblPr(tbl._tbl)
    for old in tblPr.findall(f'{{{_IA_W_NS}}}tblBorders'):
        tblPr.remove(old)
    tblBorders = etree.SubElement(tblPr, f'{{{_IA_W_NS}}}tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = etree.SubElement(tblBorders, f'{{{_IA_W_NS}}}{side}')
        el.set(f'{{{_IA_W_NS}}}val', 'none')


def _ia_set_table_full_width(tbl):
    tblPr = _ia_get_or_add_tblPr(tbl._tbl)
    tblW = tblPr.find(f'{{{_IA_W_NS}}}tblW')
    if tblW is None:
        tblW = etree.SubElement(tblPr, f'{{{_IA_W_NS}}}tblW')
    tblW.set(f'{{{_IA_W_NS}}}w', '5000')
    tblW.set(f'{{{_IA_W_NS}}}type', 'pct')


def _ia_set_row_height(row, cm_val: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for old in trPr.findall(f'{{{_IA_W_NS}}}trHeight'):
        trPr.remove(old)
    trH = etree.SubElement(trPr, f'{{{_IA_W_NS}}}trHeight')
    trH.set(f'{{{_IA_W_NS}}}val', str(int(cm_val * 567)))
    trH.set(f'{{{_IA_W_NS}}}hRule', 'atLeast')


def _ia_para_border(para, side: str, hex6: str, sz: int = 8, space: int = 1):
    """Ajoute une bordure (top/bottom/left) à un paragraphe."""
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(f'{{{_IA_W_NS}}}pBdr')
    if pBdr is None:
        pBdr = etree.SubElement(pPr, f'{{{_IA_W_NS}}}pBdr')
    for old in pBdr.findall(f'{{{_IA_W_NS}}}{side}'):
        pBdr.remove(old)
    el = etree.SubElement(pBdr, f'{{{_IA_W_NS}}}{side}')
    el.set(f'{{{_IA_W_NS}}}val', 'single')
    el.set(f'{{{_IA_W_NS}}}sz', str(sz))
    el.set(f'{{{_IA_W_NS}}}space', str(space))
    el.set(f'{{{_IA_W_NS}}}color', hex6.upper())


def _ia_insert_page_field(para_elem):
    """Insère un champ PAGE \\* MERGEFORMAT dans l'XML d'un paragraphe."""
    W = _IA_W_NS
    xml_space = '{http://www.w3.org/XML/1998/namespace}space'

    def _r_with_rpr():
        r = etree.Element(f'{{{W}}}r')
        rPr = etree.SubElement(r, f'{{{W}}}rPr')
        rf = etree.SubElement(rPr, f'{{{W}}}rFonts')
        rf.set(f'{{{W}}}ascii', 'Arial'); rf.set(f'{{{W}}}hAnsi', 'Arial')
        sz = etree.SubElement(rPr, f'{{{W}}}sz'); sz.set(f'{{{W}}}val', '16')
        szcs = etree.SubElement(rPr, f'{{{W}}}szCs'); szcs.set(f'{{{W}}}val', '16')
        clr = etree.SubElement(rPr, f'{{{W}}}color'); clr.set(f'{{{W}}}val', '808080')
        i_el = etree.SubElement(rPr, f'{{{W}}}i')
        ics = etree.SubElement(rPr, f'{{{W}}}iCs')
        return r

    r1 = _r_with_rpr(); para_elem.append(r1)
    fc1 = etree.SubElement(r1, f'{{{W}}}fldChar')
    fc1.set(f'{{{W}}}fldCharType', 'begin'); fc1.set(f'{{{W}}}dirty', '1')

    r2 = _r_with_rpr(); para_elem.append(r2)
    it = etree.SubElement(r2, f'{{{W}}}instrText')
    it.set(xml_space, 'preserve'); it.text = ' PAGE \\* MERGEFORMAT '

    r3 = _r_with_rpr(); para_elem.append(r3)
    etree.SubElement(r3, f'{{{W}}}fldChar').set(f'{{{W}}}fldCharType', 'separate')

    r4 = _r_with_rpr(); para_elem.append(r4)
    t4 = etree.SubElement(r4, f'{{{W}}}t'); t4.text = '1'

    r5 = _r_with_rpr(); para_elem.append(r5)
    etree.SubElement(r5, f'{{{W}}}fldChar').set(f'{{{W}}}fldCharType', 'end')


def _ia_setup_document(doc, devis: str):
    """Configure A4, marges, en-tête et pied de page du rapport IA."""
    W = _IA_W_NS
    sec = doc.sections[0]
    sec.page_width        = Cm(21.0)
    sec.page_height       = Cm(29.7)
    sec.left_margin       = Cm(2.5)
    sec.right_margin      = Cm(2.0)
    sec.top_margin        = Cm(2.5)
    sec.bottom_margin     = Cm(2.0)
    sec.header_distance   = Cm(1.2)
    sec.footer_distance   = Cm(1.2)
    sec.different_first_page_header_footer = True

    # ── En-tête (pages 2+) ───────────────────────────────────────────
    hdr = sec.header
    hdr.is_linked_to_previous = False
    for p in list(hdr.paragraphs):
        p._element.getparent().remove(p._element)

    p_hdr = hdr.add_paragraph()
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after  = Pt(3)
    # Tab droite à ~14 cm
    pPr_h = p_hdr._p.get_or_add_pPr()
    tabs_h = etree.SubElement(pPr_h, f'{{{W}}}tabs')
    t_h = etree.SubElement(tabs_h, f'{{{W}}}tab')
    t_h.set(f'{{{W}}}val', 'right'); t_h.set(f'{{{W}}}pos', str(int(14 * 567)))

    r_logo = p_hdr.add_run("IMMÉAU")
    r_logo.bold = True; r_logo.font.name = "Arial"; r_logo.font.size = Pt(9)
    r_logo.font.color.rgb = _ia_rgb(_IA_NAVY)

    r_tab = p_hdr.add_run("\t")
    r_tab.font.name = "Arial"; r_tab.font.size = Pt(8)

    r_ref = p_hdr.add_run(f"Rapport d'investigations — Réf. {devis}")
    r_ref.font.name = "Arial"; r_ref.font.size = Pt(8)
    r_ref.font.color.rgb = _ia_rgb('808080')

    _ia_para_border(p_hdr, 'bottom', _IA_BLUE, sz=6, space=1)

    # ── Pied de page (pages 2+) ──────────────────────────────────────
    ftr = sec.footer
    ftr.is_linked_to_previous = False
    for p in list(ftr.paragraphs):
        p._element.getparent().remove(p._element)

    p_ftr = ftr.add_paragraph()
    p_ftr.paragraph_format.space_before = Pt(3)
    p_ftr.paragraph_format.space_after  = Pt(0)
    _ia_para_border(p_ftr, 'top', _IA_BLUE, sz=6, space=1)
    # Tab droite
    pPr_f = p_ftr._p.get_or_add_pPr()
    tabs_f = etree.SubElement(pPr_f, f'{{{W}}}tabs')
    t_f = etree.SubElement(tabs_f, f'{{{W}}}tab')
    t_f.set(f'{{{W}}}val', 'right'); t_f.set(f'{{{W}}}pos', str(int(14 * 567)))

    r_ft = p_ftr.add_run(f"Rapport d'investigations {devis}")
    r_ft.italic = True; r_ft.font.name = "Arial"; r_ft.font.size = Pt(8)
    r_ft.font.color.rgb = _ia_rgb('808080')

    p_ftr.add_run("\t").font.name = "Arial"

    r_pg = p_ftr.add_run("Page ")
    r_pg.italic = True; r_pg.font.name = "Arial"; r_pg.font.size = Pt(8)
    r_pg.font.color.rgb = _ia_rgb('808080')

    _ia_insert_page_field(p_ftr._p)


def _add_cover_ia(doc, data: dict):
    """Page de garde redesignée — palette Imméau, table structurée."""
    adresse      = data.get("adresseProjet", "")
    cp           = data.get("cpProjet", "")
    ville        = data.get("villeProjet", "")
    client_raw   = data.get("client", "")
    devis        = data.get("devis", "")
    date_r       = data.get("dateRapport", "")
    redacteur    = data.get("redacteur", "")
    verificateur = data.get("verificateur", "")
    moe          = data.get("moe", "")
    cp_moe       = data.get("cpMoe", "")
    ville_moe    = data.get("villeMoe", "")
    titre_etude  = data.get("titreEtude", "") or "Investigation et diagnostic des canalisations"

    adresse_full = _build_adresse_full(adresse, cp, ville)
    client       = _strip_sdc_prefix(client_raw)
    moe_loc      = f"{moe}" + (f" — {cp_moe} {ville_moe}".strip() if (cp_moe or ville_moe) else "")

    # ── Table principale : 1 colonne, 5 lignes ──────────────────────
    tbl = doc.add_table(rows=5, cols=1)
    _ia_set_table_no_border(tbl)
    _ia_set_table_full_width(tbl)

    # ─ Ligne 0 : Bandeau entreprise (navy) ──────────────────────────
    _ia_set_row_height(tbl.rows[0], 2.8)
    c0 = tbl.rows[0].cells[0]
    _ia_set_cell_bg(c0, _IA_NAVY)

    p0a = c0.paragraphs[0]
    p0a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0a.paragraph_format.space_before = Pt(18)
    p0a.paragraph_format.space_after  = Pt(2)
    r0a = p0a.add_run("IMMÉAU")
    r0a.bold = True; r0a.font.name = "Arial"; r0a.font.size = Pt(26)
    r0a.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p0b = c0.add_paragraph()
    p0b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0b.paragraph_format.space_before = Pt(0)
    p0b.paragraph_format.space_after  = Pt(16)
    r0b = p0b.add_run("83, rue de Reuilly  ·  75012 Paris  ·  www.immeau.fr")
    r0b.font.name = "Arial"; r0b.font.size = Pt(9)
    r0b.font.color.rgb = RGBColor(0xBD, 0xD7, 0xEE)

    # ─ Ligne 1 : Titre rapport (blanc) ─────────────────────────────
    _ia_set_row_height(tbl.rows[1], 3.5)
    c1 = tbl.rows[1].cells[0]
    _ia_set_cell_bg(c1, _IA_WHITE)

    p1a = c1.paragraphs[0]
    p1a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1a.paragraph_format.space_before = Pt(28)
    p1a.paragraph_format.space_after  = Pt(4)
    r1a = p1a.add_run("RAPPORT D'INVESTIGATIONS ET DIAGNOSTIC")
    r1a.bold = True; r1a.font.name = "Arial"; r1a.font.size = Pt(17)
    r1a.font.color.rgb = _ia_rgb(_IA_NAVY)

    p1b = c1.add_paragraph()
    p1b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1b.paragraph_format.space_before = Pt(0)
    p1b.paragraph_format.space_after  = Pt(0)
    r1b = p1b.add_run("des canalisations en caves et sous-cours")
    r1b.font.name = "Arial"; r1b.font.size = Pt(12)
    r1b.font.color.rgb = _ia_rgb(_IA_BLUE)

    # ─ Ligne 2 : Adresse projet (bleu clair) ───────────────────────
    _ia_set_row_height(tbl.rows[2], 2.5)
    c2 = tbl.rows[2].cells[0]
    _ia_set_cell_bg(c2, _IA_LCYAN)

    p2a = c2.paragraphs[0]
    p2a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2a.paragraph_format.space_before = Pt(14)
    p2a.paragraph_format.space_after  = Pt(2)
    r2_sdc = p2a.add_run("SDC DU ")
    r2_sdc.font.name = "Arial"; r2_sdc.font.size = Pt(11)
    r2_sdc.font.color.rgb = _ia_rgb(_IA_NAVY)
    r2_adr = p2a.add_run(adresse_full.upper())
    r2_adr.bold = True; r2_adr.font.name = "Arial"; r2_adr.font.size = Pt(13)
    r2_adr.font.color.rgb = _ia_rgb(_IA_NAVY)

    if client:
        p2b = c2.add_paragraph()
        p2b.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2b.paragraph_format.space_before = Pt(4)
        p2b.paragraph_format.space_after  = Pt(12)
        r2b = p2b.add_run(client)
        r2b.font.name = "Arial"; r2b.font.size = Pt(10)
        r2b.font.color.rgb = _ia_rgb(_IA_BLUE)

    # ─ Ligne 3 : Informations projet (blanc) ───────────────────────
    c3 = tbl.rows[3].cells[0]
    _ia_set_cell_bg(c3, _IA_WHITE)

    info_items = [
        ("Réf. devis", devis),
        ("Date du rapport", date_r),
        ("Rédacteur", redacteur),
    ]
    if verificateur:
        info_items.append(("Vérificateur", verificateur))
    if moe:
        info_items.append(("Maître d'œuvre", moe_loc))

    p3_sp = c3.paragraphs[0]
    p3_sp.paragraph_format.space_before = Pt(14)
    p3_sp.paragraph_format.space_after  = Pt(0)

    for label, value in info_items:
        pi = c3.add_paragraph()
        pi.paragraph_format.space_before = Pt(5)
        pi.paragraph_format.space_after  = Pt(5)
        pi.paragraph_format.left_indent  = Cm(1.5)
        rl = pi.add_run(f"{label} : ")
        rl.bold = True; rl.font.name = "Arial"; rl.font.size = Pt(10)
        rl.font.color.rgb = _ia_rgb(_IA_NAVY)
        rv = pi.add_run(value or "—")
        rv.font.name = "Arial"; rv.font.size = Pt(10)

    p3_end = c3.add_paragraph()
    p3_end.paragraph_format.space_before = Pt(10)
    p3_end.paragraph_format.space_after  = Pt(0)

    # ─ Ligne 4 : Objet + pied (gris clair / navy) ──────────────────
    _ia_set_row_height(tbl.rows[4], 1.6)
    c4 = tbl.rows[4].cells[0]
    _ia_set_cell_bg(c4, _IA_NAVY)

    p4a = c4.paragraphs[0]
    p4a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4a.paragraph_format.space_before = Pt(8)
    p4a.paragraph_format.space_after  = Pt(4)
    r4a = p4a.add_run(titre_etude)
    r4a.italic = True; r4a.font.name = "Arial"; r4a.font.size = Pt(10)
    r4a.font.color.rgb = RGBColor(0xBD, 0xD7, 0xEE)

    p4b = c4.add_paragraph()
    p4b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4b.paragraph_format.space_before = Pt(0)
    p4b.paragraph_format.space_after  = Pt(8)
    r4b = p4b.add_run(f"Réf. {devis}   ·   Document confidentiel")
    r4b.font.name = "Arial"; r4b.font.size = Pt(8)
    r4b.font.color.rgb = RGBColor(0x7F, 0xA5, 0xC8)

    doc.add_page_break()


def _add_sommaire_ia(doc, has_photos: bool, has_bpe: bool):
    """Page SOMMAIRE dynamique."""
    # Titre dans un bandeau navy
    tbl_s = doc.add_table(rows=1, cols=1)
    _ia_set_table_no_border(tbl_s)
    _ia_set_table_full_width(tbl_s)
    _ia_set_row_height(tbl_s.rows[0], 1.4)
    cs = tbl_s.rows[0].cells[0]
    _ia_set_cell_bg(cs, _IA_NAVY)

    ps = cs.paragraphs[0]
    ps.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ps.paragraph_format.space_before = Pt(10)
    ps.paragraph_format.space_after  = Pt(10)
    ps.paragraph_format.left_indent  = Cm(0.5)
    rs = ps.add_run("SOMMAIRE")
    rs.bold = True; rs.font.name = "Arial"; rs.font.size = Pt(16)
    rs.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Espace
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(10)
    sp.paragraph_format.space_after  = Pt(4)

    # Liste des sections
    entries = [
        ("PRÉAMBULE", ""),
        ("I –", "Objet de la mission"),
        ("II –", "Description du site"),
        ("III –", "Géologie in situ"),
        ("IV –", "Canalisations inspectées"),
        ("V –", "Conclusions"),
    ]
    if has_photos:
        entries.append(("VI –", "Complément photographique"))
    entries.append(("VII –", "Conditions de travaux"))
    entries.append(("VIII –", "Entretien du réseau"))

    for i, (num, titre) in enumerate(entries):
        pe = doc.add_paragraph()
        pe.paragraph_format.space_before = Pt(6)
        pe.paragraph_format.space_after  = Pt(6)
        pe.paragraph_format.left_indent  = Cm(0.5)
        rn = pe.add_run(num)
        rn.bold = True; rn.font.name = "Arial"; rn.font.size = Pt(11)
        rn.font.color.rgb = _ia_rgb(_IA_NAVY)
        if titre:
            pe.add_run("  ").font.name = "Arial"
            rt = pe.add_run(titre)
            rt.font.name = "Arial"; rt.font.size = Pt(11)
        # Séparateur léger entre entrées
        if i < len(entries) - 1:
            _ia_para_border(pe, 'bottom', 'DDDDDD', sz=2, space=1)

    doc.add_page_break()


def _add_heading_ia(doc, text: str, level: int = 1):
    """Titre redesigné:
    Niveau 1 : texte navy gras + ligne bleue en bas + espacement généreux.
    Niveau 2 : barre bleue à gauche + texte bleu gras, indenté.
    """
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after  = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True; run.font.name = "Arial"; run.font.size = Pt(13)
        run.font.color.rgb = _ia_rgb(_IA_NAVY)
        _ia_para_border(p, 'bottom', _IA_BLUE, sz=6, space=2)
    else:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True; run.font.name = "Arial"; run.font.size = Pt(11)
        run.font.color.rgb = _ia_rgb(_IA_BLUE)
        _ia_para_border(p, 'left', _IA_BLUE, sz=18, space=6)
    return p


def _add_body_ia(doc, text: str, bold: bool = False):
    """Corps de texte : Arial 11, justifié."""
    if not text.strip():
        return None
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    run.font.name = "Arial"; run.font.size = Pt(11)
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(7)
    return p


def _add_bullet_ia(doc, text: str):
    """Puce redesignée : tiret bleu + texte, indenté."""
    p = doc.add_paragraph()
    r_puce = p.add_run("–  ")
    r_puce.bold = True; r_puce.font.name = "Arial"; r_puce.font.size = Pt(11)
    r_puce.font.color.rgb = _ia_rgb(_IA_BLUE)
    r_txt = p.add_run(text.strip())
    r_txt.font.name = "Arial"; r_txt.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent  = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)


def _add_photo_tables_ia(doc, photos_commentees: list):
    """Tableau photos 2 colonnes, sans titre (appelé après _add_heading_ia)."""
    if not photos_commentees:
        return
    n = len(photos_commentees)
    rows = (n + 1) // 2
    tbl = doc.add_table(rows=rows, cols=2)
    _ia_set_table_no_border(tbl)
    _ia_set_table_full_width(tbl)
    idx = 0
    for row in tbl.rows:
        for cell in row.cells:
            if idx >= n:
                # Cellule vide : fond gris clair
                _ia_set_cell_bg(cell, _IA_LGRAY)
                break
            photo = photos_commentees[idx]
            img_b64 = photo.get("image_base64", "")
            comment = photo.get("commentaire", "")
            _ia_set_cell_bg(cell, _IA_LGRAY)
            if img_b64:
                try:
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after  = Pt(4)
                    run = para.add_run()
                    run.add_picture(io.BytesIO(base64.b64decode(img_b64)), height=Cm(8.5))
                except Exception as e:
                    print(f"[WARN] Photo IA: {e}")
            if comment:
                cp = cell.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_before = Pt(2)
                cp.paragraph_format.space_after  = Pt(8)
                rc = cp.add_run(comment)
                rc.italic = True; rc.font.name = "Arial"; rc.font.size = Pt(9)
                rc.font.color.rgb = _ia_rgb('505050')
            idx += 1


def build_rapport_ia(data: dict) -> bytes:
    """
    Génère le rapport complet en utilisant Claude IA pour rédiger tout le texte dynamique.
    Assemble le résultat en Word avec python-docx.
    """
    import json as json_lib

    # ── Extraction des données ─────────────────────────────────────────
    adresse        = data.get("adresseProjet", "")
    cp             = data.get("cpProjet", "")
    ville          = data.get("villeProjet", "")
    client_raw     = data.get("client", "")
    devis          = data.get("devis", "")
    objet_mission  = data.get("objetMission", "")
    desc_site      = data.get("descriptionSite", "")
    parcelle       = data.get("parcelleCadastre", "")
    section_cad    = data.get("sectionCadastre", "")
    reglement      = data.get("reglementApplicable", "Ville de Paris")
    bpe_present    = data.get("bpePresent", False)
    bpe_phrase     = data.get("bpePhraseGeneree", "")
    batiment_zones = data.get("batimentZones", [])
    cour_zones     = data.get("courZones", [])
    reglementations = data.get("reglementationsSelectionnees", [])
    photos         = data.get("photosCommentees", [])
    photo_facade   = data.get("photoFacade", "")
    installations_items = data.get("installationsSanitairesItems", [])
    regards_noms   = data.get("regardsNonEtanchesNoms", [])
    regards_texte  = data.get("regardsNonEtanchesTexte", "")
    commentaire_colonne_ep = data.get("commentaireColonneEP", "")
    commentaire_regard = data.get("commentaireRegardLimite", "")
    fosse_prof     = data.get("fosseProfondeur", "")
    fosse_trappe   = data.get("fosseTrappe", False)
    commentaire_sep = data.get("commentaireSeparatif", "")
    commentaire_rest = data.get("commentaireRestaurants", "")
    commentaire_gar  = data.get("commentaireGarages", "")
    cond_colonne_ep  = data.get("condTravauxColonneEP", False)
    cond_terr       = data.get("condTravauxTerrassement", "")
    cond_prof       = data.get("condTravauxProfondeur", "")
    cond_sol        = data.get("condTravauxSol", "")
    cond_diff       = data.get("condTravauxDifficulte", "")
    cond_pave       = data.get("condTravauxPave", False)

    adresse_full = _build_adresse_full(adresse, cp, ville)
    client = _strip_sdc_prefix(client_raw)

    # ── Nettoyage objet_mission ────────────────────────────────────────
    objet_clean = objet_mission
    m = re.search(r'Cette mission a pour objectif\s+', objet_clean, re.IGNORECASE | re.DOTALL)
    if m:
        objet_clean = objet_clean[m.end():].strip()
    objet_clean = objet_clean.replace('\n', ' ').replace('  ', ' ').strip()

    # ── Construction du contexte pour Claude ──────────────────────────
    zones_text = ""
    for z in batiment_zones:
        zn = z.get("zoneName", "")
        if z.get("appPresentes") and z.get("appPhrase"):
            zones_text += f"\nCanalisations apparentes {zn}: {z['appPhrase']}"
        if z.get("entPresentes") and z.get("entPhrase"):
            zones_text += f"\nCanalisations enterrées {zn}: {z['entPhrase']}"
    for z in cour_zones:
        zn = z.get("zoneName", "")
        if z.get("entPresentes") and z.get("entPhrase"):
            zones_text += f"\nCanalisations enterrées espaces extérieurs ({zn}): {z['entPhrase']}"

    regl_details = ""
    if commentaire_colonne_ep:
        regl_details += f"\n- Colonne EP façade : il conviendra de mettre aux normes {commentaire_colonne_ep}"
    if commentaire_regard:
        regl_details += f"\n- Regard limite de propriété : mise en place d'un {commentaire_regard}"
    if fosse_prof:
        t = " avec trappe d'accès" if fosse_trappe else ""
        regl_details += f"\n- Ancienne fosse d'aisance : profondeur {fosse_prof}m{t}"
    if installations_items:
        regl_details += f"\n- Installations sanitaires sous-sol (clapet anti-retour) : {', '.join(installations_items)}"
    if regards_noms:
        regl_details += f"\n- Regards non étanches : {', '.join(regards_noms)}"
        if regards_texte:
            regl_details += f" — {regards_texte}"
    if commentaire_sep:
        regl_details += f"\n- Réseau séparatif : {commentaire_sep}"
    if commentaire_rest:
        regl_details += f"\n- Restaurants/commerces de bouche : {commentaire_rest}"
    if commentaire_gar:
        regl_details += f"\n- Garages/parkings : {commentaire_gar}"

    travaux_text = ""
    if cond_colonne_ep:
        travaux_text += "\n- Des travaux sur la colonne EP de façade (domaine public, DICT nécessaires)"
    if cond_terr:
        terr_label = {"manuel_total": "terrassement manuel", "mecanique": "terrassement mécanique", "mixte": "terrassement mixte (manuel et mécanique)"}.get(cond_terr, cond_terr)
        travaux_text += f"\n- Type de terrassement : {terr_label}"
    if cond_prof:
        travaux_text += f"\n- Profondeur des fouilles : 0 à {cond_prof} mètres"
    if cond_sol:
        travaux_text += f"\n- Nature du sol attendu : {cond_sol}"
    if cond_diff:
        diff_label = {"difficile": "terrassement difficile, purges possibles", "facile": "terrassement sans difficulté particulière"}.get(cond_diff, cond_diff)
        travaux_text += f"\n- Difficulté : {diff_label}"
    if cond_pave:
        travaux_text += "\n- Présence de pavés (réfection NF P98-335)"

    # ── Prompt pour Claude ─────────────────────────────────────────────
    system_prompt = (
        "Tu es un ingénieur expert en assainissement chez le bureau d'études Immeau. "
        "Tu rédiges des rapports d'investigation et diagnostic de réseaux d'assainissement. "
        "Ton style est professionnel, technique, concis, en français. "
        "Tu réponds UNIQUEMENT avec un objet JSON valide contenant les clés demandées. "
        "Pas de markdown autour du JSON, pas de texte avant ou après."
    )

    user_prompt = f"""
Génère le texte de rapport pour le projet suivant :

ADRESSE : {adresse_full}
CLIENT (SDC) : {client}
RÈGLEMENT : règlement d'assainissement de {reglement}
OBJET DE LA MISSION : {objet_clean or "inspection télévisée des réseaux EU et EP"}
DESCRIPTION DU SITE : {desc_site or f"Site situé au {adresse_full}, parcelle {parcelle} section {section_cad}"}
BPE : {"présent — " + bpe_phrase if bpe_present and bpe_phrase else "non mentionné"}
CANALISATIONS INSPECTÉES :{zones_text or " données non disponibles"}
MESURES RÉGLEMENTAIRES À PRENDRE :{regl_details or " aucune précisée"}
CONDITIONS DE TRAVAUX :{travaux_text or " données non disponibles"}

Retourne EXACTEMENT ce JSON (toutes les clés, textes en français, style rapport technique) :
{{
  "objet_mission": "1 ou 2 phrases sur l'objectif de l'étude",
  "description_site": "2-3 paragraphes décrivant le site, sa composition, ses réseaux EU/EP, séparés par \\n\\n",
  "canalisations_intro": "1-2 phrases d'introduction à la section IV",
  "conclusions_bpe": "{bpe_phrase or 'Aucun BPE mentionné dans ce rapport.'}",
  "conclusions_zones": "Paragraphes décrivant l'état de chaque zone, séparés par \\n\\n",
  "conclusions_reglements": "Paragraphes pour chaque point réglementaire, séparés par \\n\\n",
  "conditions_travaux": "2-3 paragraphes sur les conditions de terrassement et remblayage",
  "entretien": "1-2 phrases sur la préconisation d'entretien annuel"
}}
"""

    # ── Appel Claude ───────────────────────────────────────────────────
    raw = _call_claude_api(system_prompt, user_prompt, max_tokens=4000)
    # Extraire le JSON même si Claude a ajouté du markdown
    json_match = re.search(r'\{[\s\S]+\}', raw)
    if not json_match:
        raise Exception(f"Réponse Claude non valide: {raw[:200]}")
    sections = json_lib.loads(json_match.group())

    # ── Assemblage du document Word ────────────────────────────────────
    doc = Document()

    # Mise en page, marges, en-tête, pied de page
    _ia_setup_document(doc, devis)

    # ── Page de garde ──────────────────────────────────────────────────
    _add_cover_ia(doc, data)

    # ── SOMMAIRE ────────────────────────────────────────────────────────
    _add_sommaire_ia(doc, has_photos=bool(photo_facade or photos), has_bpe=bpe_present)

    # ── PRÉAMBULE (statique) ───────────────────────────────────────────
    _add_heading_ia(doc, "PRÉAMBULE", level=1)
    _add_body_ia(doc,
        "Chaque jour, vous utilisez de l'eau pour la vaisselle, la douche, la lessive, les WC… "
        "Ce sont les eaux usées domestiques qui repartent dans le réseau collectif, pour être traitées "
        "à la station d'épuration. Cette étude a pour objectif de vous fournir un rapport de synthèse "
        "des résultats obtenus lors des investigations afin de préserver le bon fonctionnement du réseau."
    )
    _add_body_ia(doc, "Notre étude s'appuie sur les principaux textes suivants :")
    _add_bullet_ia(doc, "Loi sur l'eau et les milieux aquatiques (LEMA) de décembre 2006")
    _add_bullet_ia(doc, "Le Fascicule 70 du CCTG relatif à l'exécution des travaux d'assainissement")
    _add_bullet_ia(doc, "Les normes et DTU en vigueur relatifs aux travaux d'assainissement")
    _add_bullet_ia(doc, f"Le règlement d'assainissement de {reglement}")
    doc.add_page_break()

    # ── I – OBJET DE LA MISSION ────────────────────────────────────────
    _add_heading_ia(doc, "I – OBJET DE LA MISSION", level=1)
    _add_body_ia(doc, f"La présente étude est demandée par le SDC du {adresse_full}.")
    _add_body_ia(doc, f"Cette mission a pour objectif {sections.get('objet_mission', objet_clean)}")

    # ── II – DESCRIPTION DU SITE ───────────────────────────────────────
    _add_heading_ia(doc, "II – DESCRIPTION DU SITE", level=1)
    for para_text in sections.get("description_site", desc_site or "").split("\n\n"):
        _add_body_ia(doc, para_text)

    # ── III – GÉOLOGIE IN SITU (statique) ─────────────────────────────
    _add_heading_ia(doc, "III – GÉOLOGIE IN SITU", level=1)
    _add_body_ia(doc,
        "D'un point de vue géologique, la parcelle se positionne sur les remblais reposant sur les "
        "alluvions anciennes (X/Fy), couche géologique de la région parisienne. Ces formations sont "
        "caractérisées par des remblais heterogènes pouvant dépasser 5 mètres d'épaisseur, reposant "
        "sur des sables, graviers et limons anciens. Ces données ont été confirmées par consultation "
        "de la Banque de Données du Sous-Sol (BRGM)."
    )

    # ── IV – CANALISATIONS INSPECTÉES ─────────────────────────────────
    _add_heading_ia(doc, "IV – CANALISATIONS INSPECTÉES", level=1)
    _add_body_ia(doc, sections.get("canalisations_intro",
        "Le rapport d'inspection télévisée reprenant les caractéristiques et anomalies du réseau "
        "est présenté en complément du présent rapport."
    ))
    _add_body_ia(doc,
        "Le plan de la cour et du sous-sol avec l'ensemble des réseaux inspectés et les anomalies "
        "observées est également fourni en complément."
    )
    _add_body_ia(doc, "Légende des plans :")
    _add_bullet_ia(doc, "Tracé vert : réseaux sans défauts visibles, état relativement neuf")
    _add_bullet_ia(doc, "Tracé orange : dégradations de revêtement avancées, non fuyards")
    _add_bullet_ia(doc, "Tracé rouge : très dégradés, nombreux défauts, étanchéité non garantie")

    # ── V – CONCLUSIONS ───────────────────────────────────────────────
    doc.add_page_break()
    _add_heading_ia(doc, "V – CONCLUSIONS", level=1)

    if bpe_present:
        _add_heading_ia(doc, "Branchement particulier à l'égout (BPE)", level=2)
        _add_body_ia(doc, sections.get("conclusions_bpe", bpe_phrase))

    if zones_text:
        _add_heading_ia(doc, "Canalisations inspectées", level=2)
        for para_text in sections.get("conclusions_zones", zones_text).split("\n\n"):
            _add_body_ia(doc, para_text)

    if regl_details or reglementations:
        _add_heading_ia(doc, "Points réglementaires", level=2)
        for para_text in sections.get("conclusions_reglements", regl_details).split("\n\n"):
            _add_body_ia(doc, para_text)

    # ── VI – COMPLÉMENT PHOTOGRAPHIQUE ───────────────────────────────
    if photo_facade or photos:
        doc.add_page_break()
        _add_heading_ia(doc, "VI – COMPLÉMENT PHOTOGRAPHIQUE", level=1)
        if photo_facade:
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run()
                r.add_picture(io.BytesIO(base64.b64decode(photo_facade)), height=Cm(9.87))
                _add_body_ia(doc, "Vue de la façade de l'immeuble")
            except Exception as e:
                print(f"[WARN] Photo façade IA: {e}")
        _add_photo_tables_ia(doc, photos)

    # ── VII – CONDITIONS DE TRAVAUX ───────────────────────────────────
    doc.add_page_break()
    _add_heading_ia(doc, "VII – CONDITIONS DE TRAVAUX", level=1)
    _add_heading_ia(doc, "Conditions d'accessibilité", level=2)
    _add_body_ia(doc,
        "Les travaux à réaliser sont en domaine privatif. Afin d'éventuellement pouvoir stocker "
        "des matériaux sur une place de stationnement dans la rue, une demande d'occupation "
        "temporaire du domaine public devra être déposée en mairie."
    )
    if cond_colonne_ep:
        _add_body_ia(doc,
            "D'autre part, la mise aux normes de la colonne d'eaux pluviales de façade nécessitera "
            "des travaux sur le domaine public. Dans ce cas, des DICT seront à demander aux "
            "différents concessionnaires."
        )
    for para_text in sections.get("conditions_travaux", "").split("\n\n"):
        _add_body_ia(doc, para_text)
    if cond_pave:
        _add_heading_ia(doc, "Réfection du pavage", level=2)
        _add_body_ia(doc,
            "La réfection des autobloquants devra respecter la norme française NF P98-335 de mai 2007 "
            "relative aux chaussées urbaines : Mise en œuvre des pavés et dalles en béton, des pavés "
            "en terre cuite et des pavés et dalles en pierre naturelle."
        )

    # ── VIII – ENTRETIEN DU RÉSEAU ────────────────────────────────────
    _add_heading_ia(doc, "VIII – ENTRETIEN DU RÉSEAU", level=1)
    _add_body_ia(doc,
        "Les réseaux d'assainissement et leurs ouvrages nécessitent un entretien régulier et préventif. "
        "Les canalisations entartrées ou bouchées sont source d'odeurs désagréables et peuvent entraîner "
        "des refoulements dans les parties privatives."
    )
    _add_body_ia(doc, sections.get("entretien",
        "Nous préconisons un curage des réseaux d'évacuations à raison d'une prestation par an. "
        "Pour toute demande ou renseignement, contactez-nous au 09 72 60 90 09."
    ))

    # ── Sérialise ─────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────
# Envoi du mail
# ─────────────────────────────────────────────

def send_email(docx_bytes: bytes, filename: str, devis: str, adresse: str):
    if not BREVO_API_KEY:
        raise Exception("BREVO_API_KEY manquante dans les variables d'environnement")

    body_text = (
        f"Bonjour,\n\n"
        f"Le rapport [{filename}] a été correctement généré et peut être téléchargé en pièce jointe.\n\n"
        f"Devis : {devis}\n"
        f"Adresse : {adresse}\n\n"
        f"Application Immeau."
    )

    payload = {
        "sender":      {"name": "Immeau", "email": MAIL_FROM},
        "to":          [{"email": MAIL_TO}],
        "subject":     f"Création du rapport [{filename.replace('.docx', '')}]",
        "textContent": body_text,
        "attachment":  [{"name": filename, "content": base64.b64encode(docx_bytes).decode()}],
    }

    resp = req_lib.post(
        "https://api.brevo.com/v3/smtp/email",
        headers={"api-key": BREVO_API_KEY, "Content-Type": "application/json"},
        json=payload,
        timeout=30,
    )
    if resp.status_code not in (200, 201):
        raise Exception(f"Brevo API erreur {resp.status_code}: {resp.text}")


# ─────────────────────────────────────────────
# Routes Flask
# ─────────────────────────────────────────────

@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"})


@app.route("/generer_rapport", methods=["POST"])
def generer_rapport():
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "Corps JSON manquant"}), 400

    devis    = data.get("devis", "00000")
    adresse  = data.get("adresseProjet", "")
    filename = f"Rapport d'investigations {devis}.docx"

    def process_and_send():
        try:
            docx_bytes = build_rapport(data)
            send_email(docx_bytes, filename, devis, adresse)
            print(f"[OK] Rapport {filename} envoyé à {MAIL_TO}")
        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"[ERROR] Échec génération rapport {filename} : {e}")

    thread = threading.Thread(target=process_and_send, daemon=True)
    thread.start()

    return jsonify({
        "success": True,
        "message": f"Rapport {filename} en cours de génération. Vous recevrez un e-mail à {MAIL_TO} dans quelques instants.",
    })


@app.route("/telecharger_rapport", methods=["POST"])
def telecharger_rapport():
    from flask import send_file
    try:
        data = request.get_json(force=True)
        docx_bytes = build_rapport(data)
        devis    = data.get("devis", "00000")
        filename = f"Rapport d'investigations {devis}.docx"
        return send_file(
            io.BytesIO(docx_bytes),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/generer_rapport_ia", methods=["POST"])
def generer_rapport_ia():
    data = request.get_json(force=True)
    if not data:
        return jsonify({"error": "Corps JSON manquant"}), 400

    devis    = data.get("devis", "00000")
    adresse  = data.get("adresseProjet", "")
    filename = f"Rapport IA {devis}.docx"

    def process_and_send():
        try:
            docx_bytes = build_rapport_ia(data)
            send_email(docx_bytes, filename, devis, adresse)
            print(f"[OK] Rapport IA {filename} envoyé à {MAIL_TO}")
        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"[ERROR] Échec génération rapport IA {filename} : {e}")

    thread = threading.Thread(target=process_and_send, daemon=True)
    thread.start()

    return jsonify({
        "success": True,
        "message": f"Rapport IA {filename} en cours de génération. Vous recevrez un e-mail à {MAIL_TO} dans quelques instants.",
    })


# ─────────────────────────────────────────────────────────────────────────────
# Keep-alive : empêche Render free tier de mettre le serveur en veille.
# Un thread tourne en arrière-plan et fait un GET /health toutes les 10 min.
# ─────────────────────────────────────────────────────────────────────────────
def _keepalive_loop():
    import time
    # Attendre le démarrage complet avant le premier ping
    time.sleep(60)
    health_url = os.environ.get(
        "RENDER_EXTERNAL_URL",
        "https://immeau-rapport-backend.onrender.com"
    ).rstrip("/") + "/health"
    while True:
        try:
            req_lib.get(health_url, timeout=20)
        except Exception:
            pass  # On ignore les erreurs réseau
        time.sleep(600)  # 10 minutes


_keepalive_thread = threading.Thread(target=_keepalive_loop, daemon=True)
_keepalive_thread.start()


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
